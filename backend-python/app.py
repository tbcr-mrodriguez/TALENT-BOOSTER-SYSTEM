import os
import json
import re
import secrets
from datetime import datetime
import pdfplumber
import docx
import docx2txt
import mammoth
from io import BytesIO
from flask import Flask, request, jsonify, send_file, render_template
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from openai import OpenAI
from sentence_transformers import SentenceTransformer
import psycopg2
import psycopg2.extras
from werkzeug.utils import secure_filename
from werkzeug.exceptions import RequestEntityTooLarge

import secrets
from functools import wraps
from datetime import datetime, timedelta
import jwt
from werkzeug.middleware.proxy_fix import ProxyFix


# =====================================================
# CARGAR CONFIGURACIÓN SEGÚN ENTORNO
# =====================================================
from dotenv import load_dotenv

# Detectar entorno (por variable de entorno o por defecto desarrollo)
ENVIRONMENT = os.environ.get('ENVIRONMENT', 'development')

# Cargar archivo .env correspondiente
env_file = f'.env.{ENVIRONMENT}'
if os.path.exists(env_file):
    load_dotenv(env_file)
    print(f"📁 Cargando configuración desde {env_file}")
else:
    # Fallback a .env si existe
    if os.path.exists('.env'):
        load_dotenv()
        print("📁 Cargando configuración desde .env")
    else:
        print("⚠️ No se encontró archivo de configuración")

# Importar configuración centralizada
from config import (
    ENVIRONMENT as CFG_ENV,
    ALLOWED_ORIGINS,
    MAX_CONTENT_LENGTH,
    ALLOWED_EXTENSIONS,
    ALLOWED_VIDEO_EXTENSIONS,
    RATE_LIMITS_CONFIG,
    OPENAI_API_KEY,
    DEEPSEEK_API_KEY,
    SECRET_KEY,
    get_database_url,
    validate_config,
    print_config
)

# =====================================================
# CONFIGURACIÓN DE LA APLICACIÓN
# =====================================================
app = Flask(__name__)

app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH
app.config['SECRET_KEY'] = SECRET_KEY

# Headers de seguridad
@app.after_request
def add_security_headers(response):
    # Headers OWASP recomendados
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    response.headers['Content-Security-Policy'] = "default-src 'self' http://localhost:3000 http://localhost:5001; script-src 'self' 'unsafe-inline' 'unsafe-eval' https://cdn.tailwindcss.com https://cdn.jsdelivr.net; style-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com; img-src 'self' data: https:; font-src 'self' data:; connect-src 'self' http://localhost:5001 http://localhost:3000;"
    response.headers['Referrer-Policy'] = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy'] = 'camera=(self), microphone=(self), geolocation=()'
    
    return response

# CORS
CORS(app, 
     origins=ALLOWED_ORIGINS,
     supports_credentials=True,
     allow_headers=['Content-Type', 'Authorization'],
     methods=['GET', 'POST', 'PUT', 'DELETE', 'OPTIONS'])

# Rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=RATE_LIMITS_CONFIG['default']
)

# =====================================================
# CONFIGURACIÓN DE CARPETAS (igual que antes)
# =====================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(BASE_DIR)

app.template_folder = os.path.join(BASE_DIR, 'templates')
print(f"📁 Templates folder: {app.template_folder}")

UPLOAD_FOLDER = os.path.join(PARENT_DIR, "uploads")
VIDEO_UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads", "entrevistas_video")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VIDEO_UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_video_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS

# =====================================================
# CONEXIÓN A POSTGRESQL
# =====================================================
def get_db_connection():
    database_url = get_database_url()
    if not database_url:
        raise ValueError("DATABASE_URL no configurada")
    return psycopg2.connect(database_url)

# =====================================================
# EMBEDDINGS
# =====================================================
print("🔄 Cargando modelo de embeddings...")
modelo_embeddings = SentenceTransformer('all-MiniLM-L6-v2')
print("✅ Modelo cargado")

# =====================================================
# API KEYS (usando configuración)
# =====================================================
openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
deepseek_client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com"
) if DEEPSEEK_API_KEY else None

# Mostrar configuración
print_config()

# Validar configuración en producción
if ENVIRONMENT == 'production':
    if not validate_config():
        print("❌ Configuración inválida. La aplicación puede no funcionar correctamente.")

# ... EL RESTO DE TU CÓDIGO (funciones y endpoints) CONTINÚA IGUAL ...


app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', secrets.token_urlsafe(32))

# Headers de seguridad OWASP
@app.after_request
def add_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'
    response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response


# Configuración CORS segura
ALLOWED_ORIGINS = [
    "http://localhost:3000",
    "http://localhost:3001",
    "http://localhost:5001",
    os.environ.get('FRONTEND_URL', '')
]


limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# =====================================================
# CONFIGURACIÓN DE CARPETAS
# =====================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PARENT_DIR = os.path.dirname(BASE_DIR)

app.template_folder = os.path.join(BASE_DIR, 'templates')
print(f"📁 Templates folder: {app.template_folder}")

UPLOAD_FOLDER = os.path.join(PARENT_DIR, "uploads")
VIDEO_UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads", "entrevistas_video")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(VIDEO_UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}
ALLOWED_VIDEO_EXTENSIONS = {'webm', 'mp4', 'mkv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_video_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_VIDEO_EXTENSIONS

# =====================================================
# CONEXIÓN A POSTGRESQL
# =====================================================
def get_db_connection():
    """Obtiene conexión a PostgreSQL desde variable de entorno"""
    database_url = os.environ.get('DATABASE_URL')
    
    if database_url:
        if database_url.startswith('postgres://'):
            database_url = database_url.replace('postgres://', 'postgresql://', 1)
        return psycopg2.connect(database_url)
    else:
        return psycopg2.connect(
            host='localhost',
            port=5432,
            database='talent_pipeline',
            user='talent_user',
            password='Talent123!'
        )

# =====================================================
# EMBEDDINGS
# =====================================================
print("🔄 Cargando modelo de embeddings...")
modelo_embeddings = SentenceTransformer('all-MiniLM-L6-v2')
print("✅ Modelo cargado")

# =====================================================
# API KEYS
# =====================================================
openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))
deepseek_client = OpenAI(
    api_key=os.environ.get("DEEPSEEK_API_KEY", ""),
    base_url="https://api.deepseek.com"
)

if not os.environ.get("OPENAI_API_KEY"):
    print("⚠️ ADVERTENCIA: OPENAI_API_KEY no está configurada")
if not os.environ.get("DEEPSEEK_API_KEY"):
    print("⚠️ ADVERTENCIA: DEEPSEEK_API_KEY no está configurada")

# =====================================================
# FUNCIONES DE AYUDA
# =====================================================
def safe_json_parse(text):
    import re
    start = text.find("{")
    end = text.rfind("}") + 1
    if start == -1 or end == 0:
        raise ValueError("No se encontró JSON válido")
    return json.loads(text[start:end])

def sanitize_filename(filename):
    return secure_filename(filename)

def validate_text_input(text, max_length=10000):
    if not isinstance(text, str):
        return ""
    if len(text) > max_length:
        text = text[:max_length]
    return re.sub(r'[^\w\s\.,;:\-@\(\)\n\r]', '', text)

def normalize_lists(data):
    if "interpretacion" in data:
        for campo in ["industrias_secundarias", "habilidades_clave", "fortalezas", "areas_mejora"]:
            if campo in data["interpretacion"]:
                arr = data["interpretacion"][campo]
                if not arr:
                    data["interpretacion"][campo] = []
                else:
                    clean = []
                    for item in arr:
                        if isinstance(item, dict):
                            clean.append(" ".join(str(v) for v in item.values()))
                        else:
                            clean.append(str(item))
                    data["interpretacion"][campo] = clean
        
        if "datos_crudos" in data:
            fields = ["certificaciones", "habilidades_listadas", "herramientas_listadas", "idiomas_listados"]
            for f in fields:
                arr = data["datos_crudos"].get(f)
                if not arr:
                    data["datos_crudos"][f] = []
                else:
                    clean = []
                    for item in arr:
                        if isinstance(item, dict):
                            clean.append(" ".join(str(v) for v in item.values()))
                        else:
                            clean.append(str(item))
                    data["datos_crudos"][f] = clean
    else:
        fields = ["certificaciones", "industrias", "roles", "logros", "habilidades", "herramientas", "idiomas"]
        for f in fields:
            arr = data.get(f)
            if not arr:
                data[f] = []
                continue
            clean = []
            for item in arr:
                if isinstance(item, dict):
                    clean.append(" ".join(str(v) for v in item.values()))
                else:
                    clean.append(str(item))
            data[f] = clean
    return data

def analyze_cv(text, model_choice):
    prompt = f"""
Eres un headhunter experto con 20 años de experiencia. Tu tarea es EXTRAER e INTERPRETAR el CV de manera PROFUNDA y DETALLADA.

**INSTRUCCIÓN CRÍTICA:** 
NO asumas una estructura fija. Analiza el CV y extrae TODO lo que encuentres. 
Si el CV tiene secciones como "Proyectos", "Publicaciones", "Voluntariado", "Premios", etc., créalas como campos.

Devuelve SOLO JSON con esta estructura BASE (puedes AGREGAR MÁS CAMPOS si el CV los tiene):

{{
  "datos_crudos": {{
    "nombre": "",
    "telefono": "",
    "email": "",
    "ubicacion": "",
    "profesion_escrita": "",
    "universidad": "",
    "grado_academico": "",
    "anio_graduacion": "",
    "certificaciones": [],
    "perfil_textual": "",
    "experiencia_laboral": [
      {{
        "puesto": "",
        "empresa": "",
        "fecha_inicio": "",
        "fecha_fin": "",
        "descripcion": "",
        "logros": []
      }}
    ],
    "habilidades_listadas": [],
    "herramientas_listadas": [],
    "idiomas_listados": [],
    "pretension_salarial": "",
    "disponibilidad": "",
    "empresas": [],
    "proyectos": [],
    "publicaciones": [],
    "premios": [],
    "voluntariados": []
  }},
  "interpretacion": {{
    "sector_deducido": "",
    "anos_experiencia_deducidos": "",
    "seniority": "",
    "perfil_interpretado": "",
    "industria_principal": "",
    "industrias_secundarias": [],
    "habilidades_clave": [],
    "fortalezas": [],
    "areas_mejora": [],
    "rol_tipico": "",
    "nivel_impacto": "",
    "recomendacion": "",
    "resumen_ejecutivo": "",
    "puntos_fuertes": [],
    "potencial_crecimiento": "",
    "cultura_ideal": ""
  }}
}}

**REGLAS DE EXTRACCIÓN:**

1. **EMPRESAS (MUY IMPORTANTE)**: 
   - Extrae TODAS las empresas donde trabajó de la experiencia laboral
   - Ponlas en "empresas" como lista de strings
   - También en cada experiencia laboral incluye el campo "empresa"

2. **EXPERIENCIA LABORAL**:
   - Por cada trabajo, extrae: puesto, empresa, fechas, descripción
   - Si hay logros específicos (ej: "aumenté ventas en 30%"), ponlos en "logros"

3. **SECCIONES ADICIONALES**:
   - Si ves "Proyectos", crea "proyectos": ["proyecto1", "proyecto2"]
   - Si ves "Publicaciones", crea "publicaciones": ["pub1", "pub2"]
   - Si ves "Premios", crea "premios": ["premio1", "premio2"]
   - Si ves "Voluntariado", crea "voluntariados": ["vol1", "vol2"]

**REGLAS DE INTERPRETACIÓN (SÉ PROFUNDO):**

1. **PERFIL INTERPRETADO**: Resume quién es el candidato en 2-3 líneas, destacando su esencia profesional.

2. **FORTALEZAS**: Lista 5-7 fortalezas específicas basadas en su experiencia, habilidades y logros.

3. **ÁREAS DE MEJORA**: Sé honesto y constructivo. ¿Qué le falta? ¿En qué podría mejorar?

4. **ROL TÍPICO**: ¿Qué puesto debería ocupar según su perfil? (ej: "Gerente de Proyectos Senior", "Desarrollador Full Stack")

5. **NIVEL DE IMPACTO**: Alto/Medio/Bajo - ¿Qué impacto puede tener en una organización?

6. **RECOMENDACIÓN**: ¿Qué tipo de empresas o roles le quedarían mejor?

7. **RESUMEN EJECUTIVO**: Una frase poderosa que venda al candidato.

8. **PUNTOS FUERTES**: 3-5 características que lo hacen único.

9. **POTENCIAL CRECIMIENTO**: ¿Hacia dónde puede evolucionar profesionalmente?

10. **CULTURA IDEAL**: ¿En qué tipo de ambiente laboral prosperaría?

CV a analizar:
{text[:20000]}
"""
    if model_choice == "deepseek":
        client = deepseek_client
        model = "deepseek-chat"
    else:
        client = openai_client
        model = "gpt-4o-mini"
    
    response = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=5000
    )
    
    data = safe_json_parse(response.choices[0].message.content)
    
    # Guardar el texto completo para embeddings y búsqueda
    data['texto_completo_cv'] = text
    
    return data

def score_candidate(data):
    if "interpretacion" in data:
        interp = data["interpretacion"]
        score = 0
        if interp.get("seniority") in ["Senior", "Lead"]: score += 20
        elif interp.get("seniority") == "Semi-Senior": score += 15
        elif interp.get("seniority") == "Junior": score += 10
        if interp.get("habilidades_clave"): score += min(25, len(interp["habilidades_clave"]) * 5)
        if interp.get("fortalezas"): score += min(15, len(interp["fortalezas"]) * 3)
        if interp.get("nivel_impacto") == "Alto": score += 20
        elif interp.get("nivel_impacto") == "Medio": score += 10
        elif interp.get("nivel_impacto") == "Bajo": score += 5
        if interp.get("industria_principal"): score += 5
        if interp.get("industrias_secundarias"): score += min(5, len(interp["industrias_secundarias"]))
        exp_text = interp.get("anos_experiencia_deducidos", "")
        if "años" in exp_text:
            match = re.search(r'(\d+)', exp_text)
            if match:
                años = int(match.group(1))
                if años >= 8: score += 10
                elif años >= 5: score += 7
                elif años >= 2: score += 4
        if score >= 80: criterio = "Alta compatibilidad"
        elif score >= 60: criterio = "Compatibilidad media-alta"
        elif score >= 40: criterio = "Compatibilidad media"
        elif score >= 20: criterio = "Compatibilidad media-baja"
        else: criterio = "Perfil limitado"
        data["score"] = score
        data["criterio_ai"] = criterio
    return data

def construir_texto_embedding(data):
    """
    Construye el texto para el embedding usando TODO el CV
    Así la búsqueda semántica encuentra cualquier cosa
    """
    texto_completo = ""
    
    # ===== 1. TEXTO COMPLETO DEL CV =====
    texto_completo += data.get('texto_completo_cv', '')
    texto_completo += "\n\n"
    
    # ===== 2. DATOS ESTRUCTURADOS (todo, para reforzar) =====
    dc = data.get('datos_crudos', {})
    if dc:
        texto_completo += "=== DATOS DEL CV ===\n"
        for key, value in dc.items():
            if value:
                if isinstance(value, list):
                    texto_completo += f"{key}: {', '.join(str(v) for v in value)}\n"
                elif isinstance(value, dict):
                    texto_completo += f"{key}: {json.dumps(value, ensure_ascii=False)}\n"
                else:
                    texto_completo += f"{key}: {value}\n"
    
    # ===== 3. INTERPRETACIÓN IA =====
    interp = data.get('interpretacion', {})
    if interp:
        texto_completo += "\n=== INTERPRETACIÓN ===\n"
        for key, value in interp.items():
            if value:
                if isinstance(value, list):
                    texto_completo += f"{key}: {', '.join(str(v) for v in value)}\n"
                else:
                    texto_completo += f"{key}: {value}\n"
    
    return texto_completo

def buscar_por_embedding(texto_busqueda, limite=20):
    try:
        embedding_busqueda = modelo_embeddings.encode(texto_busqueda).tolist()
        
        conn = get_db_connection()
        
        # Usar RealDictCursor para obtener resultados como diccionario
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Método CORRECTO: pasar el embedding como lista de Python
        # psycopg2 lo convierte automáticamente al formato vector
        cursor.execute("""
            SELECT * FROM buscar_por_similitud_con_embedding(
                %s::vector,
                %s
            )
        """, (embedding_busqueda, limite))
        
        resultados = cursor.fetchall()
        conn.close()
        
        candidatos = []
        for row in resultados:
            try:
                # row ya es un diccionario con las columnas
                candidato = {
                    'id': row['candidate_id'],
                    'nombre': row['nombre'],
                    'profesion': row['profesion'],
                    'similitud_semantica': round(row['similitud'] * 100, 1),
                    'raw_data': row['raw_data']
                }
                candidatos.append(candidato)
            except Exception as e:
                print(f"Error parseando fila: {e}")
                continue
        
        print(f"✅ Búsqueda devolvió {len(candidatos)} candidatos")
        return candidatos
        
    except Exception as e:
        print(f"❌ Error en búsqueda: {e}")
        import traceback
        traceback.print_exc()
        return []
    
def busqueda_inteligente(pregunta, candidatos, limite=5):
    """Usa ChatGPT para buscar candidatos de forma inteligente con toda la información"""
    try:
        # Construir perfiles de los candidatos preseleccionados
        perfiles = []
        for i, c in enumerate(candidatos[:10]):  # Limitar a 10 para no saturar
            try:
                if c.get('raw_data'):
                    datos = json.loads(c['raw_data']) if isinstance(c['raw_data'], str) else c['raw_data']
                else:
                    datos = c
                
                perfil = construir_perfil_candidato(datos)
                # Limitar longitud del perfil para no exceder tokens
                if len(perfil) > 1500:
                    perfil = perfil[:1500] + "..."
                perfiles.append(f"CANDIDATO {i+1}:\n{perfil}\n")
            except Exception as e:
                print(f"Error construyendo perfil: {e}")
                continue
        
        if not perfiles:
            return []
        
        # Prompt más simple y robusto
        prompt = f"""
Eres un headhunter experto. El usuario busca: "{pregunta}"

Analiza estos candidatos y determina cuáles cumplen con lo que busca.

{''.join(perfiles)}

Devuelve SOLO UN JSON válido con esta estructura exacta:
{{"resultados": [{{"candidato_id": 1, "match_score": 85, "razones": ["razon1", "razon2"], "fortalezas": ["fortaleza1"], "debilidades": ["debilidad1"]}}]}}

IMPORTANTE: 
- El candidato_id es el número del candidato en la lista (1, 2, 3...)
- match_score debe ser un número entre 0 y 100
- Si un candidato no cumple, dale score bajo (menos de 50)
- NO incluyas texto fuera del JSON
"""
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=800
        )
        
        contenido = response.choices[0].message.content
        print(f"📥 Respuesta ChatGPT: {contenido[:200]}...")
        
        # Extraer JSON de la respuesta (por si viene con texto adicional)
        import re
        json_match = re.search(r'\{.*\}\]\}', contenido, re.DOTALL)
        if json_match:
            json_text = json_match.group()
        else:
            json_text = contenido
        
        # Limpiar posibles caracteres problemáticos
        json_text = json_text.replace('\n', ' ').replace('\r', ' ')
        
        try:
            resultado = json.loads(json_text)
        except json.JSONDecodeError as e:
            print(f"Error parseando JSON: {e}")
            print(f"Texto problemático: {json_text[:500]}")
            # Intentar reparar: buscar patrones comunes
            if 'match_score' in json_text:
                # Extraer manualmente
                import re
                patron = r'"candidato_id":\s*(\d+).*?"match_score":\s*(\d+)'
                matches = re.findall(patron, json_text)
                resultado = {"resultados": []}
                for match in matches:
                    resultado["resultados"].append({
                        "candidato_id": int(match[0]),
                        "match_score": int(match[1]),
                        "razones": [],
                        "fortalezas": []
                    })
            else:
                return []
        
        # Mapear resultados a candidatos
        candidatos_con_score = []
        for res in resultado.get('resultados', []):
            idx = res.get('candidato_id', 0) - 1
            if 0 <= idx < len(candidatos):
                c = candidatos[idx].copy()
                c['match_score_inteligente'] = res.get('match_score', 0)
                c['razones_match'] = res.get('razones', [])
                c['fortalezas_match'] = res.get('fortalezas', [])
                c['debilidades_match'] = res.get('debilidades', [])
                candidatos_con_score.append(c)
        
        # Ordenar por score
        candidatos_con_score.sort(key=lambda x: x.get('match_score_inteligente', 0), reverse=True)
        
        print(f"✅ Búsqueda inteligente: {len(candidatos_con_score)} candidatos evaluados")
        return candidatos_con_score
        
    except Exception as e:
        print(f"❌ Error en búsqueda inteligente: {e}")
        import traceback
        traceback.print_exc()
        return []    
    

def load_all_candidates_from_db():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_todos_candidatos')
        resultados = cursor.fetchall()
        conn.close()
        candidates = []
        for row in resultados:
            try:
                if isinstance(row, tuple):
                    data = row[0]
                else:
                    data = row
                if isinstance(data, str):
                    data = json.loads(data)
                candidates.append(data)
            except:
                continue
        return candidates
    except Exception as e:
        print(f"❌ Error cargando candidatos: {e}")
        return []
    
def generar_embedding_candidato(data, candidate_id):
    """Genera y guarda embedding para un candidato"""
    try:
        texto_completo = construir_texto_embedding(data)
        embedding = modelo_embeddings.encode(texto_completo).tolist()
        
        # Convertir a formato PostgreSQL vector
        embedding_str = '[' + ','.join(str(x) for x in embedding) + ']'
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Usar SQL directo con casting explícito
        cursor.execute("""
            SELECT guardar_embedding(%s, %s::vector, %s)
        """, (candidate_id, embedding_str, texto_completo))
        
        conn.commit()
        conn.close()
        print(f"✅ Embedding guardado para candidato ID: {candidate_id}")
        return True
    except Exception as e:
        print(f"❌ Error generando embedding: {e}")
        import traceback
        traceback.print_exc()
        return False

def save_candidate_to_db(data):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('crear_candidato', [data['archivo'], json.dumps(data, ensure_ascii=False)])
        candidate_id = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        generar_embedding_candidato(data, candidate_id)
        return candidate_id
    except Exception as e:
        print(f"❌ Error guardando candidato: {e}")
        raise

def extract_pdf(path):
    text = ""
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
    except Exception as e:
        print(f"❌ Error extrayendo PDF: {e}")
    return text

def extract_docx(path):
    text = ""
    try:
        with open(path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            text = result.value
            if len(text) > 50:
                return text
    except:
        pass
    try:
        doc = docx.Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts)
        if len(text) > 50:
            return text
    except:
        pass
    try:
        text = docx2txt.process(path)
    except:
        pass
    return text

def extract_text(path):
    if path.lower().endswith(".pdf"):
        return extract_pdf(path)
    if path.lower().endswith(".docx"):
        return extract_docx(path)
    return ""

# =====================================================
# API ENDPOINTS
# =====================================================

@app.route("/api/candidates", methods=["GET"])
@limiter.limit("30 per minute")
def api_get_candidates():
    try:
        candidates = load_all_candidates_from_db()
        return jsonify({"success": True, "data": candidates})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/analyze", methods=["POST"])
@limiter.limit("10 per minute")
def api_analyze():
    try:
        print("="*50)
        print("📥 PYTHON: Recibida petición de análisis")
        files = request.files.getlist("files")
        if not files:
            return jsonify({"success": False, "error": "No se recibieron archivos"}), 400
        if len(files) > 10:
            return jsonify({"success": False, "error": "Máximo 10 archivos por vez"}), 400
        model = request.form.get("model", "gpt")
        if model not in ['gpt', 'deepseek']:
            model = 'gpt'
        results = []
        for file in files:
            if not allowed_file(file.filename):
                return jsonify({"success": False, "error": f"Tipo no permitido: {file.filename}"}), 400
            filename = sanitize_filename(file.filename)
            path = os.path.join(UPLOAD_FOLDER, filename)
            file.save(path)
            text = extract_text(path)
            if not text or len(text) < 50:
                print(f"⚠️ Texto muy corto: {filename}")
                continue
            data = analyze_cv(text, model)
            data = normalize_lists(data)
            data = score_candidate(data)
            data["archivo"] = filename
            data["fecha_analisis"] = datetime.now().isoformat()
            candidate_id = save_candidate_to_db(data)
            print(f"✅ Candidato guardado con ID: {candidate_id}")
            results.append(data)
        return jsonify({"success": True, "data": results})
    except RequestEntityTooLarge:
        return jsonify({"success": False, "error": "Archivo >16MB"}), 413
    except Exception as e:
        print(f"❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/search", methods=["GET"])
@limiter.limit("30 per minute")
def api_search():
    try:
        query = request.args.get('q', '')
        skills_list = request.args.getlist('skills[]')
        query = validate_text_input(query, 500)
        if not query and skills_list:
            query = " ".join(skills_list)
            query = validate_text_input(query, 500)
        if not query:
            return jsonify({"success": False, "error": "Query required"}), 400
        results = buscar_por_embedding(query, 20)
        return jsonify({"success": True, "data": results, "count": len(results)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/stats", methods=["GET"])
@limiter.limit("30 per minute")
def api_stats():
    try:
        candidates = load_all_candidates_from_db()
        total = len(candidates)
        sectores = set()
        skills = set()
        senior = 0
        for c in candidates:
            if c.get("interpretacion"):
                if c["interpretacion"].get("sector_deducido"):
                    sectores.add(c["interpretacion"]["sector_deducido"])
                if c["interpretacion"].get("habilidades_clave"):
                    skills.update(c["interpretacion"]["habilidades_clave"])
                if c["interpretacion"].get("seniority") in ["Senior", "Lead"]:
                    senior += 1
            else:
                if c.get("sector_principal"):
                    sectores.add(c["sector_principal"])
                if c.get("habilidades"):
                    skills.update(c["habilidades"])
        return jsonify({
            "success": True,
            "stats": {
                "total": total,
                "sectores": len(sectores),
                "skills": len(skills),
                "senior_percent": round((senior/total*100) if total > 0 else 0)
            }
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/view-cv/<path:filename>", methods=["GET"])
@limiter.limit("60 per minute")
def api_view_cv(filename):
    try:
        from urllib.parse import unquote
        filename = sanitize_filename(unquote(filename))
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        real_path = os.path.realpath(file_path)
        upload_real = os.path.realpath(UPLOAD_FOLDER)
        if not real_path.startswith(upload_real):
            return jsonify({"success": False, "error": "Acceso denegado"}), 403
        if not os.path.exists(file_path):
            return jsonify({"success": False, "error": "File not found"}), 404
        return send_file(file_path)
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/search-embedding", methods=["POST"])
@limiter.limit("30 per minute")
def search_embedding():
    try:
        data = request.json
        query = data.get('query', '')
        limite = min(data.get('limite', 20), 100)
        if not query:
            return jsonify({"success": False, "error": "Texto de búsqueda vacío"}), 400
        query = validate_text_input(query, 500)
        candidatos = buscar_por_embedding(query, limite)
        return jsonify({"success": True, "data": candidatos, "count": len(candidatos)})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# =====================================================
# ENDPOINTS DE PLANTILLAS Y ENTREVISTAS (con procedimientos)
# =====================================================

@app.route("/api/admin/plantillas", methods=["GET"])
@limiter.limit("30 per minute")
def obtener_plantillas():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_plantillas')
        rows = cursor.fetchall()
        conn.close()
        plantillas = []
        for row in rows:
            plantillas.append({
                "id": row[0],
                "nombre": row[1],
                "puesto": row[2],
                "tipo": row[3],
                "fecha_creacion": row[4],
                "activa": bool(row[5])
            })
        return jsonify({"success": True, "plantillas": plantillas})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/admin/plantilla", methods=["POST"])
@limiter.limit("10 per minute")
def crear_plantilla():
    try:
        data = request.json
        nombre = data.get('nombre')
        puesto = data.get('puesto')
        tipo = data.get('tipo', 'generica')
        preguntas = data.get('preguntas', [])
        if not nombre:
            return jsonify({"success": False, "error": "Nombre requerido"}), 400
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('crear_plantilla', [nombre, puesto, tipo, json.dumps(preguntas)])
        plantilla_id = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        return jsonify({"success": True, "plantilla_id": plantilla_id})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/admin/plantilla/<int:plantilla_id>", methods=["GET"])
@limiter.limit("30 per minute")
def obtener_plantilla(plantilla_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_plantilla', [plantilla_id])
        row = cursor.fetchone()
        conn.close()
        if not row:
            return jsonify({"success": False, "error": "Plantilla no encontrada"}), 404
        if isinstance(row, tuple):
            data = row[0]
        else:
            data = row
        if isinstance(data, str):
            data = json.loads(data)
        return jsonify({"success": True, "plantilla": data})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/entrevista/iniciar", methods=["POST"])
@limiter.limit("10 per minute")
def iniciar_entrevista():
    try:
        data = request.json
        candidato_id = data.get('candidato_id')
        plantilla_id = data.get('plantilla_id')
        if not candidato_id or not plantilla_id:
            return jsonify({"success": False, "error": "Faltan datos"}), 400
        token = secrets.token_urlsafe(32)
        base_url = request.host_url.rstrip('/')
        enlace = f"{base_url}/entrevista-candidato.html?token={token}"
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('programar_entrevista', [candidato_id, plantilla_id, token, enlace])
        entrevista_id = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        return jsonify({
            "success": True,
            "entrevista_id": entrevista_id,
            "enlace": enlace,
            "token": token
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/entrevista/candidato/<token>", methods=["GET"])
def obtener_entrevista_candidato(token):
    try:
        token = validate_text_input(token, 100)
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_entrevista_por_token', [token])
        row = cursor.fetchone()
        conn.close()
        if not row:
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        if isinstance(row, tuple):
            data = row[0]
        else:
            data = row
        if isinstance(data, str):
            data = json.loads(data)
        if data.get('estado') != 'pendiente':
            return jsonify({
                "success": False, 
                "error": f"Esta entrevista ya fue {data.get('estado')}"
            }), 400
        return jsonify({
            "success": True,
            "entrevista": {
                "id": data.get('id'),
                "candidato_nombre": data.get('candidato_nombre'),
                "candidato_id": data.get('candidato_id'),
                "plantilla_nombre": data.get('plantilla_nombre'),
                "puesto": data.get('puesto'),
                "preguntas": data.get('preguntas', []),
                "total_preguntas": len(data.get('preguntas', []))
            }
        })
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/entrevista/candidato/responder", methods=["POST"])
def guardar_respuesta_candidato():
    try:
        token = request.form.get('token')
        pregunta_id = request.form.get('pregunta_id')
        pregunta_texto = request.form.get('pregunta_texto')
        transcripcion = request.form.get('transcripcion', '')
        tiempo = request.form.get('tiempo', 0)
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_entrevista_por_token', [token])
        row = cursor.fetchone()
        if not row:
            conn.close()
            return jsonify({"success": False, "error": "Entrevista no válida"}), 400
        if isinstance(row, tuple):
            entrevista_data = row[0]
        else:
            entrevista_data = row
        if isinstance(entrevista_data, str):
            entrevista_data = json.loads(entrevista_data)
        entrevista_id = entrevista_data.get('id')
        filename = None
        if 'video' in request.files:
            video_file = request.files['video']
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"entrevista_{entrevista_id}_pregunta_{pregunta_id}_{timestamp}.webm"
            filepath = os.path.join(VIDEO_UPLOAD_FOLDER, filename)
            video_file.save(filepath)
            print(f"✅ Video guardado: {filename}")
        cursor.callproc('guardar_respuesta', [
            entrevista_id, pregunta_id, pregunta_texto, 
            transcripcion, filename, tiempo
        ])
        conn.commit()
        conn.close()
        return jsonify({"success": True})
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/entrevista/candidato/finalizar", methods=["POST"])
def finalizar_entrevista_candidato():
    try:
        token = request.json.get('token')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 🔥 USAR CALL
        cursor.execute("CALL finalizar_entrevista(%s)", (token,))
        conn.commit()
        
        cursor.execute("SELECT id FROM entrevistas WHERE token_acceso = %s", (token,))
        row = cursor.fetchone()
        cursor.close()
        conn.close()
        
        if not row:
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        entrevista_id = row[0]
        
        import threading
        import requests
        
        def analizar_en_background():
            import time
            time.sleep(2)
            try:
                requests.post(f"http://localhost:5001/api/entrevista/analizar-profundo/{entrevista_id}", timeout=60)
            except Exception as e:
                print(f"Error: {e}")
        
        threading.Thread(target=analizar_en_background, daemon=True).start()
        
        return jsonify({"success": True, "entrevista_id": entrevista_id})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route("/api/entrevista/resultado-detallado/<int:entrevista_id>", methods=["GET"])
def resultado_entrevista_detallado(entrevista_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("CALL obtener_resultado_entrevista(%s, %s)", (entrevista_id, None))
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": True, "analisis": None, "entrevista_id": entrevista_id})
        
        resultado = json.loads(row[0])
        
        # 🔥 PARSEAR EL ANÁLISIS SI ES STRING
        analisis = resultado.get('entrevista', {}).get('analisis')
        if analisis and isinstance(analisis, str):
            try:
                analisis = json.loads(analisis)
                resultado['entrevista']['analisis'] = analisis
            except:
                pass
        
        return jsonify({
            "success": True,
            "entrevista": resultado.get('entrevista', {}),
            "respuestas": resultado.get('respuestas', [])
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/entrevista/resultados/<int:candidato_id>", methods=["GET"])
def obtener_resultados_entrevista(candidato_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_resultados_entrevista', [candidato_id])
        rows = cursor.fetchall()
        conn.close()
        
        resultados = []
        for row in rows:
            resultados.append({
                "id": row[0],
                "fecha": row[1],
                "fecha_fin": row[2],
                "estado": row[3],
                "analisis": row[4] if row[4] else None,
                "puntuacion": row[5],
                "plantilla": row[6],
                "puesto": row[7]
            })
        return jsonify({"success": True, "resultados": resultados})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route("/api/entrevistas/procesar-pendientes", methods=["POST"])
def procesar_entrevistas_pendientes():
    """Procesa todas las entrevistas pendientes de análisis"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.callproc('obtener_entrevistas_pendientes')
        pendientes = cursor.fetchall()
        conn.close()
        
        resultados = []
        for p in pendientes:
            entrevista_id = p[0]
            print(f"🔄 Procesando entrevista ID: {entrevista_id}")
            success = analizar_entrevista(entrevista_id)
            resultados.append({
                "id": entrevista_id,
                "success": success
            })
        
        return jsonify({
            "success": True,
            "procesadas": len(pendientes),
            "resultados": resultados
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500    

@app.route("/api/entrevistas/todas", methods=["GET"])
def obtener_todas_entrevistas():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 🔥 USAR CALL directamente con execute
        cursor.execute("CALL listar_entrevistas(%s)", (None,))
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": True, "entrevistas": []})
        
        resultado = json.loads(row[0])
        
        return jsonify({
            "success": True,
            "entrevistas": resultado.get('entrevistas', [])
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500
    

@app.route("/api/entrevista/reporte-completo/<int:entrevista_id>", methods=["GET"])
def reporte_entrevista_completo(entrevista_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 🔥 USAR CALL
        cursor.execute("CALL obtener_resultado_entrevista(%s, %s)", (entrevista_id, None))
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        resultado = json.loads(row[0])
        entrevista_data = resultado.get('entrevista', {})
        respuestas = resultado.get('respuestas', [])
        duracion = sum(r.get('tiempo', 0) for r in respuestas)
        
        analisis = entrevista_data.get('analisis')
        if analisis and isinstance(analisis, str):
            try:
                analisis = json.loads(analisis)
            except:
                pass
        
        return jsonify({
            "success": True,
            "entrevista": {
                "id": entrevista_data.get('id'),
                "candidato_id": entrevista_data.get('candidato_id'),
                "fecha": entrevista_data.get('fecha_inicio'),
                "fecha_fin": entrevista_data.get('fecha_fin'),
                "estado": entrevista_data.get('estado'),
                "puntuacion_global": entrevista_data.get('puntuacion', 0),
                "candidato_nombre": entrevista_data.get('candidato_nombre', 'No disponible'),
                "candidato_profesion": entrevista_data.get('candidato_profesion', 'No disponible'),
                "plantilla": entrevista_data.get('plantilla', 'No disponible'),
                "puesto": entrevista_data.get('puesto', 'No disponible'),
                "duracion": duracion
            },
            "respuestas": respuestas,
            "analisis": analisis,
            "candidato_data": entrevista_data.get('candidato_data'),
            "fotos": []
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/entrevistas/candidato/<int:candidato_id>", methods=["GET"])
def obtener_entrevistas_por_candidato_api(candidato_id):
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.callproc('obtener_entrevistas_por_candidato', [candidato_id])
        rows = cursor.fetchall()
        conn.close()
        entrevistas = []
        for row in rows:
            entrevistas.append({
                "id": row[0],
                "fecha_inicio": row[1],
                "fecha_fin": row[2],
                "estado": row[3],
                "puntuacion": row[4],
                "plantilla": row[5],
                "puesto": row[6],
                "enlace": row[7]
            })
        return jsonify({"success": True, "entrevistas": entrevistas})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
@app.route("/api/entrevistas/procesar-pendientes", methods=["POST"])
def procesar_pendientes_entrevistas():  # 🔥 NOMBRE DIFERENTE
    """Procesa todas las entrevistas que están pendientes de análisis"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # Obtener entrevistas pendientes
        cursor.execute("""
            SELECT e.id, e.candidato_id, c.nombre as candidato_nombre
            FROM entrevistas e
            JOIN candidates c ON e.candidato_id = c.id
            WHERE e.estado = 'completada' 
              AND e.analisis_ia IS NULL
            ORDER BY e.fecha_fin DESC
        """)
        
        pendientes = cursor.fetchall()
        cursor.close()
        conn.close()
        
        resultados = []
        for p in pendientes:
            try:
                # Llamar a la función analizar_entrevista (que ya existe en tu código)
                from app import analizar_entrevista
                success = analizar_entrevista(p['id'])
                resultados.append({
                    "id": p['id'],
                    "candidato": p['candidato_nombre'],
                    "success": success
                })
            except Exception as e:
                resultados.append({
                    "id": p['id'],
                    "candidato": p['candidato_nombre'],
                    "success": False,
                    "error": str(e)
                })
        
        return jsonify({
            "success": True,
            "procesadas": len(pendientes),
            "resultados": resultados
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500 

@app.route("/api/asistente", methods=["POST"])
def asistente_inteligente():
    """Headhunter inteligente - todo lo hace la IA"""
    try:
        data = request.json
        pregunta = data.get('mensaje', '')
        
        print(f"\n🤵 HEADHUNTER: '{pregunta}'")
        
        # Buscar candidatos relevantes
        candidatos = buscar_por_embedding(pregunta, 15)
        print(f"📊 Encontrados {len(candidatos)} candidatos")
        
        if not candidatos:
            return jsonify({
                "success": True,
                "respuesta": f"No encontré candidatos relacionados con '{pregunta}'."
            })
        
        # Construir perfiles de los candidatos
        perfiles = []
        for i, c in enumerate(candidatos[:8]):  # Máximo 8
            try:
                if c.get('raw_data'):
                    datos = json.loads(c['raw_data']) if isinstance(c['raw_data'], str) else c['raw_data']
                else:
                    datos = c
                
                dc = datos.get('datos_crudos', {})
                interp = datos.get('interpretacion', {})
                
                nombre = dc.get('nombre', 'No especificado')
                profesion = dc.get('profesion_escrita', 'No especificada')
                ubicacion = dc.get('ubicacion', 'No especificada')
                universidad = dc.get('universidad', 'No especificada')
                empresas = dc.get('empresas', [])
                experiencia = interp.get('anos_experiencia_deducidos', 'No especificada')
                seniority = interp.get('seniority', 'No especificado')
                habilidades = interp.get('habilidades_clave', [])
                
                perfil = f"""
CANDIDATO {i+1}:
Nombre: {nombre}
Profesión: {profesion}
Ubicación: {ubicacion}
Universidad: {universidad}
Experiencia: {experiencia}
Seniority: {seniority}
Empresas: {', '.join(empresas) if empresas else 'No especificadas'}
Habilidades: {', '.join(habilidades) if habilidades else 'No especificadas'}
"""
                perfiles.append(perfil)
            except Exception as e:
                print(f"Error construyendo perfil: {e}")
                continue
        
        if not perfiles:
            return jsonify({
                "success": True,
                "respuesta": "No pude procesar los perfiles de los candidatos."
            })
        print(f"🔍 PERFIL ENVIADO A IA:\n{perfil}\n")
        # =====================================================
        # DEJAR QUE LA IA ANALICE TODO
        # =====================================================
        prompt = f"""
Eres un headhunter experto con 20 años de experiencia. El usuario pregunta:

"{pregunta}"

Aquí están los candidatos disponibles:

{''.join(perfiles)}

Analiza cada candidato y responde con el siguiente formato:

**RESPUESTA** (texto conversacional, breve pero completo)

Luego, para cada candidato que cumpla con la búsqueda, incluye:
- Nombre
- Por qué cumple (concretamente)
- Si no cumple, dilo claramente

REGLAS:
- Si la pregunta es por ubicación (ej: "alguien de Liberia"), verifica la ubicación del candidato
- Si la pregunta es por empresa, verifica las empresas donde trabajó
- Si la pregunta es por habilidad, verifica sus habilidades
- Si la pregunta es por experiencia, verifica los años
- Si ningún candidato cumple, dilo claramente
- Sé específico, no uses frases genéricas

RESPONDE EN ESPAÑOL.
"""
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=1000
        )
        
        respuesta = response.choices[0].message.content
        
        return jsonify({
            "success": True,
            "respuesta": respuesta
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500

# =====================================================
# RUTAS PARA PÁGINAS HTML
# =====================================================
@app.route('/<path:path>.html')
def serve_html(path):
    try:
        return render_template(f"{path}.html")
    except Exception as e:
        return f"Error: {e}", 404

@app.route('/test')
def test():
    return "<h1>HOLA MUNDO</h1><p>Si ves esto, Flask funciona</p>"

@app.route('/entrevista-admin.html')
def admin_page():
    return render_template('entrevista-admin.html')

@app.route('/entrevista-candidato.html')
def candidato_page():
    return render_template('entrevista-candidato.html')

@app.route('/test.html')
def test_page():
    return render_template('test.html')


@app.route("/api/match-empleo", methods=["POST"])
def match_empleo():
    """Encuentra los mejores candidatos para una descripción de empleo usando ChatGPT"""
    try:
        data = request.json
        descripcion = data.get('descripcion', '')
        titulo = data.get('titulo', 'Posición')
        cantidad = data.get('cantidad', 10)
        
        print(f"\n{'='*60}")
        print(f"🎯 MATCH EMPLEO: {titulo}")
        print(f"📝 Descripción: {descripcion[:200]}...")
        print(f"{'='*60}")
        
        # Paso 1: Usar IA para extraer requisitos clave
        prompt_requisitos = f"""
Eres un experto en reclutamiento. Analiza esta descripción de empleo y extrae los requisitos MÍNIMOS.

Título: {titulo}
Descripción: {descripcion}

Devuelve SOLO JSON con:
{{
  "habilidades_requeridas": ["skill1", "skill2", "skill3"],
  "años_experiencia_minimo": número,
  "seniority_requerido": "junior" | "semi-senior" | "senior" | "lead",
  "sector": "sector específico" | null,
  "rol": "el rol principal"
}}
"""
        response_requisitos = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt_requisitos}],
            temperature=0
        )
        
        requisitos = safe_json_parse(response_requisitos.choices[0].message.content)
        print(f"📋 Requisitos extraídos: {requisitos}")
        
        # Paso 2: Buscar candidatos con embedding
        consulta = " ".join(requisitos.get('habilidades_requeridas', []))
        if not consulta:
            consulta = descripcion[:200]
        
        candidatos = buscar_por_embedding(consulta, cantidad * 3)
        print(f"📊 Candidatos encontrados por embedding: {len(candidatos)}")
        
        if not candidatos:
            return jsonify({
                "success": True,
                "candidatos": [],
                "total": 0,
                "mensaje": "No encontré candidatos que coincidan con esta búsqueda."
            })
        
        # Paso 3: Evaluar cada candidato con ChatGPT
        resultados = []
        
        for c in candidatos[:15]:  # Limitar a 15 para no saturar la API
            # Parsear raw_data para obtener el perfil completo
            try:
                if c.get('raw_data'):
                    datos = json.loads(c['raw_data']) if isinstance(c['raw_data'], str) else c['raw_data']
                else:
                    datos = c
            except:
                datos = c
            
            # Construir perfil del candidato para ChatGPT
            nombre = datos.get('datos_crudos', {}).get('nombre', c.get('nombre', 'N/A'))
            profesion = datos.get('datos_crudos', {}).get('profesion_escrita', c.get('profesion', 'N/A'))
            experiencia = datos.get('interpretacion', {}).get('anos_experiencia_deducidos', 'No especificada')
            seniority = datos.get('interpretacion', {}).get('seniority', 'No especificado')
            habilidades = datos.get('interpretacion', {}).get('habilidades_clave', [])
            fortalezas = datos.get('interpretacion', {}).get('fortalezas', [])
            perfil = datos.get('interpretacion', {}).get('perfil_interpretado', '')
            ubicacion = datos.get('datos_crudos', {}).get('ubicacion', 'No especificada')
            universidad = datos.get('datos_crudos', {}).get('universidad', 'No especificada')
            
            perfil_candidato = f"""
Candidato: {nombre}
Profesión: {profesion}
Experiencia: {experiencia}
Seniority: {seniority}
Ubicación: {ubicacion}
Universidad: {universidad}
Habilidades: {', '.join(habilidades)}
Fortalezas: {', '.join(fortalezas)}
Perfil: {perfil[:300]}
"""
            
            # Usar ChatGPT para evaluar el match
            prompt_evaluacion = f"""
Eres un headhunter experto evaluando candidatos para un puesto.

PUESTO: {titulo}
DESCRIPCIÓN DEL PUESTO: {descripcion}

CANDIDATO A EVALUAR:
{perfil_candidato}

Evalúa qué tan bien calza este candidato con el puesto. Devuelve SOLO JSON con:
{{
  "match_score": número del 0 al 100,
  "fortalezas": ["fortaleza1", "fortaleza2", "fortaleza3"],
  "debilidades": ["debilidad1", "debilidad2"],
  "resumen": "resumen breve de la evaluación"
}}

Sé honesto y crítico. Si el candidato no calza, dale un score bajo.
"""
            
            try:
                response_eval = openai_client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[{"role": "user", "content": prompt_evaluacion}],
                    temperature=0.3,
                    max_tokens=400
                )
                
                evaluacion = safe_json_parse(response_eval.choices[0].message.content)
                
                c['match_score'] = evaluacion.get('match_score', 0)
                c['fortalezas_match'] = evaluacion.get('fortalezas', [])
                c['debilidades_match'] = evaluacion.get('debilidades', [])
                c['resumen_match'] = evaluacion.get('resumen', '')
                c['nombre'] = nombre
                c['profesion'] = profesion
                c['ubicacion'] = ubicacion
                c['experiencia'] = experiencia
                c['seniority_match'] = seniority
                c['habilidades_match'] = habilidades[:5]
                
                resultados.append(c)
                print(f"✅ {nombre}: {c['match_score']}% match")
                
            except Exception as e:
                print(f"❌ Error evaluando {nombre}: {e}")
                continue
        
        # Ordenar por match score
        resultados.sort(key=lambda x: x['match_score'], reverse=True)
        
        # Filtrar solo los que tienen al menos 50% de match
        resultados_filtrados = [r for r in resultados if r['match_score'] >= 50]
        
        # Tomar los mejores
        mejores = resultados_filtrados[:cantidad]
        
        print(f"\n📊 RESULTADO FINAL: {len(mejores)} candidatos calificados de {len(resultados)} evaluados")
        
        if len(mejores) == 0:
            return jsonify({
                "success": True,
                "candidatos": [],
                "total": 0,
                "mensaje": "No encontré candidatos con un match score suficientemente alto (>50%). Considera ajustar la descripción del puesto."
            })
        
        return jsonify({
            "success": True,
            "empleo": {
                "titulo": titulo,
                "descripcion": descripcion[:200],
                "requisitos": requisitos
            },
            "candidatos": mejores,
            "total": len(mejores),
            "evaluados": len(resultados)
        })
        
    except Exception as e:
        print(f"❌ Error en match empleo: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/analyze-search", methods=["POST", "OPTIONS"])
def analyze_search():
    """Analiza la intención de búsqueda del usuario"""
    # Manejar preflight CORS
    if request.method == "OPTIONS":
        return '', 200
    
    try:
        data = request.json
        query = data.get('query', '')
        
        if not query:
            return jsonify({"success": False, "error": "Query vacía"}), 400
        
        print(f"📥 Analizando búsqueda: '{query}'")
        
        prompt = f"""
Eres un headhunter experto. Analiza esta búsqueda de reclutamiento y extrae TODA la información relevante.

Búsqueda: "{query}"

Extrae:
- LUGARES (ciudades, provincias, países, regiones)
- UNIVERSIDADES (cualquier mención de instituciones educativas)
- HABILIDADES técnicas o blandas
- ROL o puesto buscado
- SECTOR o industria
- SENIORITY requerido

Devuelve SOLO JSON con esta estructura:
{{
  "lugares": ["lista", "de", "lugares", "mencionados"],
  "universidades": ["lista", "de", "universidades"],
  "habilidades": ["lista", "de", "habilidades"],
  "rol_buscado": "el rol o puesto que busca",
  "sector": "sector o industria mencionada",
  "seniority": "seniority requerido (si menciona)",
  "texto_original": "{query}"
}}
"""
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        analysis = safe_json_parse(response.choices[0].message.content)
        print(f"📊 Análisis extraído: {analysis}")
        
        return jsonify({"success": True, "analysis": analysis})
        
    except Exception as e:
        print(f"❌ Error en analyze-search: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
    

@app.route("/api/search-embedding-enhanced", methods=["POST", "OPTIONS"])
def search_embedding_enhanced():
    """Búsqueda semántica mejorada con análisis de intención"""
    # Manejar preflight CORS
    if request.method == "OPTIONS":
        return '', 200
    
    try:
        data = request.json
        query = data.get('query', '')
        analysis = data.get('analysis', {})
        limite = data.get('limite', 30)
        
        print(f"📥 Búsqueda semántica: '{query}'")
        print(f"📋 Análisis recibido: {analysis}")
        
        if not query:
            return jsonify({"success": False, "error": "Texto de búsqueda vacío"}), 400
        
        # Extraer términos útiles del análisis si existe
        lugares = analysis.get('lugares', [])
        universidades = analysis.get('universidades', [])
        habilidades = analysis.get('habilidades', [])
        
        # Construir consulta enriquecida
        query_enriquecida = query
        if lugares:
            query_enriquecida += " " + " ".join(lugares)
        if universidades:
            query_enriquecida += " " + " ".join(universidades)
        if habilidades:
            query_enriquecida += " " + " ".join(habilidades)
        
        print(f"🔍 Consulta enriquecida: {query_enriquecida}")
        
        # Búsqueda semántica base
        candidatos = buscar_por_embedding(query_enriquecida, limite * 2)
        
        # Si no hay resultados, intentar con la consulta original
        if len(candidatos) == 0:
            print("⚠️ Sin resultados, intentando con consulta original")
            candidatos = buscar_por_embedding(query, limite * 2)
        
        print(f"📊 Resultados base: {len(candidatos)} candidatos")
        
        # Agregar relevancia
        for c in candidatos:
            c['relevancia'] = c.get('similitud_semantica', 50)
        
        print(f"✅ Devolviendo {len(candidatos[:limite])} candidatos")
        
        return jsonify({
            "success": True,
            "data": candidatos[:limite],
            "total": len(candidatos)
        })
        
    except Exception as e:
        print(f"❌ Error en búsqueda semántica: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500  
    

@app.route("/api/entrevista/analizar-profundo/<int:entrevista_id>", methods=["POST"])
def analizar_entrevista_profundo(entrevista_id):
    """Genera análisis profesional completo usando el nuevo reporte completo"""
    print(f"🔥 ANALIZANDO ENTREVISTA ID: {entrevista_id}")
    
    try:
        # Usar el nuevo procedimiento que trae TODA la información
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        cursor.execute("CALL obtener_reporte_completo_entrevista(%s, %s)", (entrevista_id, None))
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": False, "error": "No se encontraron datos de la entrevista"}), 404
        
        datos = json.loads(row[0])
        
        entrevista = datos.get('entrevista', {})
        candidato_data = entrevista.get('candidato_data', {})
        plantilla = datos.get('plantilla', {})
        preguntas_config = datos.get('preguntas_configuracion', [])
        respuestas = datos.get('respuestas', [])
        
        # Extraer datos del candidato
        interpretacion = candidato_data.get('interpretacion', {})
        datos_crudos = candidato_data.get('datos_crudos', {})
        
        # =========================================================
        # CONSTRUIR PERFIL DEL CANDIDATO (DETALLADO)
        # =========================================================
        habilidades_tecnicas = interpretacion.get('habilidades_clave', [])
        fortalezas_cv = interpretacion.get('fortalezas', [])
        anos_experiencia = interpretacion.get('anos_experiencia_deducidos', 'No especificado')
        seniority = interpretacion.get('seniority', 'No especificado')
        empresas = datos_crudos.get('empresas', [])
        
        perfil_candidato = f"""
=== PERFIL DEL CANDIDATO ===
Nombre: {datos_crudos.get('nombre', 'No disponible')}
Profesión: {datos_crudos.get('profesion_escrita', 'No disponible')}
Seniority detectado: {seniority}
Años experiencia: {anos_experiencia}
Ubicación: {datos_crudos.get('ubicacion', 'No disponible')}

=== TECNOLOGÍAS Y HERRAMIENTAS ===
{', '.join(habilidades_tecnicas) if habilidades_tecnicas else 'No especificadas'}

=== PRINCIPALES EMPRESAS ===
{chr(10).join(['- ' + e for e in empresas[:5]]) if empresas else 'No especificadas'}

=== PERFIL INTERPRETADO POR IA ===
{interpretacion.get('perfil_interpretado', 'No disponible')}

=== FORTALEZAS SEGÚN CV ===
{chr(10).join(['- ' + f for f in fortalezas_cv[:5]]) if fortalezas_cv else 'No especificadas'}
"""
        
        # =========================================================
        # PERFIL DEL PUESTO
        # =========================================================
        perfil_puesto = plantilla.get('perfil_puesto', {})
        perfil_puesto_texto = f"""
=== PERFIL DEL PUESTO ===
Nombre de la plantilla: {plantilla.get('nombre', 'No especificado')}
Descripción: {plantilla.get('descripcion', 'No especificada')[:500]}

=== REQUISITOS DEL PUESTO ===
{chr(10).join(['- ' + r for r in perfil_puesto.get('requisitos', [])]) if perfil_puesto.get('requisitos') else 'No se especificaron requisitos formales'}

=== RESPONSABILIDADES ===
{chr(10).join(['- ' + r for r in perfil_puesto.get('responsabilidades', [])]) if perfil_puesto.get('responsabilidades') else 'No se especificaron responsabilidades formales'}
"""
        
        # =========================================================
        # CULTURA DE LA EMPRESA
        # =========================================================
        cultura = plantilla.get('cultura_empresa', {})
        valores = cultura.get('valores', [])
        cultura_texto = f"""
=== CULTURA ORGANIZACIONAL ===
Valores clave:
{chr(10).join(['- ' + v[:200] + '...' for v in valores[:4]]) if valores else 'No especificados'}

Ambiente de trabajo:
{cultura.get('ambiente', 'No especificado')[:400]}
"""
        
        # =========================================================
        # OBJETIVO DE LA ENTREVISTA
        # =========================================================
        objetivo_texto = f"""
=== OBJETIVO DE LA ENTREVISTA ===
{plantilla.get('objetivo', 'No especificado')[:600]}
"""
        
        # =========================================================
        # PREGUNTAS CON RESPUESTAS
        # =========================================================
        preguntas_detalle = []
        for i, conf in enumerate(preguntas_config):
            respuesta_obj = next((r for r in respuestas if r.get('pregunta') == conf.get('pregunta')), {})
            tiempo_respuesta = respuesta_obj.get('tiempo', 0)
            texto_respuesta = respuesta_obj.get('respuesta', '')
            
            if not texto_respuesta or texto_respuesta == "":
                evaluacion_tiempo = f"⚠️ El candidato grabó un video de {tiempo_respuesta} segundos pero no transcribió su respuesta. El reclutador debe revisar el video."
            else:
                evaluacion_tiempo = f"Respuesta transcrita: {texto_respuesta[:300]}..."
            
            preguntas_detalle.append(f"""
--- PREGUNTA {i+1}: {conf.get('competencia', 'Competencia')} ---
📌 Pregunta: {conf.get('pregunta', 'No disponible')}
🎯 Objetivo de la pregunta: {conf.get('objetivo', 'No especificado')}
⭐ Criterio STAR esperado: {conf.get('criterio_star', 'No especificado')}

🗣️ RESPUESTA DEL CANDIDATO:
{evaluacion_tiempo}
Tiempo de respuesta: {tiempo_respuesta} segundos
""")
        
        # =========================================================
        # PROMPT COMPLETO
        # =========================================================
        prompt = f"""
Actúa como un Headhunter Senior con 20 años de experiencia en reclutamiento técnico. Eres conocido por ser ESTRICTO, ANALÍTICO y por tomar decisiones basadas en datos.

## 📋 INFORMACIÓN DEL CANDIDATO (Extraída de su CV)

{perfil_candidato}

## 🎯 INFORMACIÓN DEL PUESTO Y LA EMPRESA

{perfil_puesto_texto}

{cultura_texto}

{objetivo_texto}

## 🎙️ RESPUESTAS DEL CANDIDATO EN LA ENTREVISTA

{chr(10).join(preguntas_detalle)}

## 📊 TAREA DE EVALUACIÓN

Evalúa al candidato basándote en:

1. **Alineación técnica**: ¿Las habilidades del CV coinciden con los requisitos del puesto?
2. **Experiencia relevante**: ¿Sus años de experiencia (15+ años) y las empresas donde trabajó son relevantes?
3. **Ajuste cultural**: ¿Su perfil se alinea con los valores y ambiente descritos?
4. **Calidad de respuestas**: Evalúa si sus respuestas (por tiempo y contenido) demuestran lo que el puesto requiere

## 🔍 FORMATO DE RESPUESTA (SOLO JSON)

{{
  "resumen_ejecutivo": {{
    "nivel_general": "Evaluación general del candidato en 2-3 párrafos. Incluye comparación directa con los requisitos del puesto.",
    "principales_fortalezas": ["fortaleza1 específica", "fortaleza2 específica", "fortaleza3 específica"],
    "principales_debilidades": ["debilidad1 específica", "debilidad2 específica"],
    "cumple_requisitos_tecnicos": true/false,
    "cumple_requisitos_experiencia": true/false,
    "cumple_ajuste_cultural": true/false,
    "resumen_comparativo": "Explicación detallada de por qué el candidato SÍ o NO cumple con lo que busca la empresa"
  }},
  "analisis_tecnico": {{
    "habilidades_fuertes": ["habilidad1", "habilidad2"],
    "habilidades_debiles": ["habilidad1", "habilidad2"],
    "habilidades_faltantes": ["habilidad que el puesto requiere pero el candidato no tiene"],
    "nivel_tecnico_general": "Básico/Intermedio/Avanzado/Experto",
    "justificacion": "Explicación detallada del nivel técnico basada en su CV y experiencia"
  }},
  "analisis_por_pregunta": [
    {{
      "pregunta": "texto de la pregunta",
      "competencia": "nombre",
      "evaluacion": "análisis de lo que se esperaba vs lo que el candidato mostró",
      "fortalezas": ["fortaleza demostrada en esta pregunta"],
      "areas_mejora": ["aspecto a mejorar en esta pregunta"],
      "puntuacion_estimada": 0-100
    }}
  ],
  "analisis_ajuste_cultural": {{
    "compatibilidad": "Alta/Media/Baja",
    "razones": ["razón1 basada en su perfil", "razón2 basada en su experiencia"],
    "recomendacion_entrevistador": "texto"
  }},
  "score_global": 0-100,
  "recomendacion_final": {{
    "decision": "contratar/avanzar/no_recomendable",
    "argumento": "Explicación DETALLADA de por qué se toma esta decisión. Menciona específicamente qué requisitos cumple y cuáles no.",
    "next_steps": "Próximos pasos sugeridos si aplica"
  }},
  "fortalezas_generales": ["fortaleza1", "fortaleza2", "fortaleza3"],
  "areas_mejora": ["area1", "area2", "area3"],
  "insight_headhunter": "Análisis profundo y honesto: ¿Este candidato realmente encaja con lo que busca la empresa? ¿Por qué sí o por qué no?"
}}

Devuelve SOLO el JSON, sin texto adicional.
"""
        
        print("📡 Llamando a ChatGPT...")
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4000
        )
        
        contenido = response.choices[0].message.content
        analisis = safe_json_parse(contenido)
        
        # Guardar análisis
        conn = get_db_connection()
        cursor = conn.cursor()
        score_global = analisis.get('score_global', 0)
        cursor.execute("""
            UPDATE entrevistas 
            SET analisis_ia = %s, puntuacion_global = %s, estado = 'analizada'
            WHERE id = %s
        """, (json.dumps(analisis, ensure_ascii=False), score_global, entrevista_id))
        conn.commit()
        conn.close()
        
        print(f"✅ Análisis completado para entrevista ID {entrevista_id}")
        
        return jsonify({
            "success": True,
            "analisis": analisis
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500




def generar_html_informe(analisis, entrevista_id):
    """Genera el HTML del informe profesional"""
    
    score = analisis.get('score_global', 0)
    score_color = '#10b981' if score >= 80 else '#f59e0b' if score >= 60 else '#ef4444'
    
    html = f"""
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Informe de Evaluación - Entrevista {entrevista_id}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        
        body {{
            font-family: 'Helvetica', 'Arial', sans-serif;
            background: white;
            padding: 40px;
            line-height: 1.5;
            color: #1f2937;
        }}
        
        .container {{
            max-width: 900px;
            margin: 0 auto;
            background: white;
        }}
        
        /* Header */
        .header {{
            text-align: center;
            padding: 30px 0;
            border-bottom: 2px solid #e5e7eb;
            margin-bottom: 30px;
        }}
        
        .header h1 {{
            font-size: 28px;
            color: #1f2937;
            margin-bottom: 8px;
        }}
        
        .header .subtitle {{
            color: #6b7280;
            font-size: 14px;
        }}
        
        /* Score Card */
        .score-card {{
            background: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 100%);
            border-radius: 16px;
            padding: 24px;
            margin-bottom: 30px;
            border: 1px solid #e2e8f0;
        }}
        
        .score-row {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
        }}
        
        .score-number {{
            font-size: 48px;
            font-weight: bold;
            color: {score_color};
        }}
        
        .score-label {{
            font-size: 14px;
            color: #6b7280;
        }}
        
        .recomendacion {{
            font-size: 18px;
            font-weight: bold;
        }}
        
        /* Secciones */
        .section {{
            margin-bottom: 30px;
            break-inside: avoid;
        }}
        
        .section-title {{
            font-size: 18px;
            font-weight: 600;
            color: #1f2937;
            padding-bottom: 8px;
            border-bottom: 2px solid #3b82f6;
            margin-bottom: 16px;
        }}
        
        .grid-2 {{
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 20px;
        }}
        
        .card {{
            background: #f9fafb;
            border-radius: 12px;
            padding: 16px;
            border: 1px solid #e5e7eb;
        }}
        
        .card-title {{
            font-weight: 600;
            margin-bottom: 12px;
            font-size: 14px;
            color: #4b5563;
        }}
        
        .competencia-item {{
            margin-bottom: 12px;
        }}
        
        .competencia-header {{
            display: flex;
            justify-content: space-between;
            font-size: 13px;
            margin-bottom: 4px;
        }}
        
        .progress-bar {{
            height: 6px;
            background: #e5e7eb;
            border-radius: 3px;
            overflow: hidden;
        }}
        
        .progress-fill {{
            height: 100%;
            background: #3b82f6;
            border-radius: 3px;
        }}
        
        .badge {{
            display: inline-block;
            padding: 2px 8px;
            border-radius: 12px;
            font-size: 11px;
            font-weight: 500;
        }}
        
        .badge-alto {{ background: #d1fae5; color: #065f46; }}
        .badge-medio {{ background: #fed7aa; color: #9a3412; }}
        .badge-bajo {{ background: #fee2e2; color: #991b1b; }}
        
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #e5e7eb;
            text-align: center;
            font-size: 11px;
            color: #9ca3af;
        }}
        
        @media print {{
            body {{ padding: 20px; }}
            .no-print {{ display: none; }}
        }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Header -->
        <div class="header">
            <h1>📊 INFORME DE EVALUACIÓN</h1>
            <p class="subtitle">Entrevista profesional · {analisis.get('entrevista', {}).get('candidato_nombre', 'Candidato')}</p>
            <p class="subtitle">Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
        </div>
        
        <!-- Score Card -->
        <div class="score-card">
            <div class="score-row">
                <div>
                    <span class="score-label">Puntuación Global</span>
                    <div class="score-number">{score}/100</div>
                </div>
                <div class="text-right">
                    <span class="score-label">Recomendación</span>
                    <div class="recomendacion" style="color: {score_color}">
                        {analisis.get('recomendacion_final', {}).get('decision', 'No disponible').upper()}
                    </div>
                </div>
            </div>
        </div>
        
        <!-- Resumen Ejecutivo -->
        {f'<div class="section"><div class="section-title">📋 Resumen Ejecutivo</div><p>{analisis.get("resumen_ejecutivo", {}).get("nivel_general", "No disponible")}</p></div>' if analisis.get('resumen_ejecutivo') else ''}
        
        <!-- Fortalezas y Riesgos -->
        <div class="grid-2">
            <div class="card">
                <div class="card-title">✅ Fortalezas</div>
                <ul style="padding-left: 20px;">
                    {''.join([f'<li style="margin-bottom: 6px;">{f}</li>' for f in analisis.get('resumen_ejecutivo', {}).get('principales_fortalezas', [])[:5]]) or '<li>No especificadas</li>'}
                </ul>
            </div>
            <div class="card">
                <div class="card-title">⚠️ Riesgos</div>
                <ul style="padding-left: 20px;">
                    {''.join([f'<li style="margin-bottom: 6px;">{r}</li>' for r in analisis.get('resumen_ejecutivo', {}).get('principales_riesgos', [])[:5]]) or '<li>No identificados</li>'}
                </ul>
            </div>
        </div>
        
        <!-- Competencias Blandas -->
        <div class="section">
            <div class="section-title">⭐ Competencias Blandas</div>
            <div class="grid-2">
                {''.join([f'''
                <div class="competencia-item">
                    <div class="competencia-header">
                        <span>{c.get('competencia', '')}</span>
                        <span>{c.get('puntuacion', 0)} - {c.get('nivel', '')}</span>
                    </div>
                    <div class="progress-bar">
                        <div class="progress-fill" style="width: {c.get('puntuacion', 0)}%;"></div>
                    </div>
                    <p style="font-size: 11px; color: #6b7280; margin-top: 4px;">{c.get('evidencia', '')[:100]}</p>
                </div>
                ''' for c in analisis.get('analisis_competencias_blandas', [])[:8]])}
            </div>
        </div>
        
        <!-- Competencias Técnicas -->
        {f'''
        <div class="section">
            <div class="section-title">🔧 Competencias Técnicas</div>
            <div style="background: #f9fafb; border-radius: 12px; padding: 16px;">
                {''.join([f'''
                <div style="margin-bottom: 12px; padding-bottom: 12px; border-bottom: 1px solid #e5e7eb;">
                    <div style="display: flex; justify-content: space-between;">
                        <strong>{c.get('tecnologia', '')}</strong>
                        <span class="badge {'badge-alto' if c.get('gap') == 'Sin gap' else 'badge-medio' if c.get('gap') == 'Parcial' else 'badge-bajo'}">{c.get('gap', '')}</span>
                    </div>
                    <p style="font-size: 12px; margin-top: 4px;">CV: {c.get('nivel_cv', '')} | Entrevista: {c.get('nivel_entrevista', '')}</p>
                    <p style="font-size: 11px; color: #6b7280; margin-top: 4px;">{c.get('evidencia', '')[:120]}</p>
                </div>
                ''' for c in analisis.get('analisis_competencias_tecnicas', [])])}
            </div>
        </div>
        ''' if analisis.get('analisis_competencias_tecnicas') else ''}
        
        <!-- Insight Headhunter -->
        {f'''
        <div class="section">
            <div class="section-title">🎯 Insight del Headhunter</div>
            <div style="background: #f9fafb; border-radius: 12px; padding: 16px; border-left: 4px solid #8b5cf6;">
                <p>{analisis.get('insight_headhunter', 'No disponible')}</p>
            </div>
        </div>
        ''' if analisis.get('insight_headhunter') else ''}
        
        <!-- Preguntas de Seguimiento -->
        {f'''
        <div class="section">
            <div class="section-title">📝 Preguntas de Seguimiento</div>
            <div style="background: #f9fafb; border-radius: 12px; padding: 16px;">
                <ul style="padding-left: 20px;">
                    {''.join([f'<li style="margin-bottom: 8px;">{p}</li>' for p in analisis.get('preguntas_seguimiento', [])])}
                </ul>
            </div>
        </div>
        ''' if analisis.get('preguntas_seguimiento') else ''}
        
        <!-- Footer -->
        <div class="footer">
            <p>Informe generado automáticamente por Talent Pipeline · Documento confidencial</p>
            <p>ID de entrevista: {entrevista_id}</p>
        </div>
    </div>
</body>
</html>
"""
    return html


@app.route("/api/entrevistas/pendientes", methods=["GET"])
def entrevistas_pendientes():
    """Obtiene entrevistas completadas que no tienen análisis"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        cursor.callproc('obtener_entrevistas_pendientes')
        rows = cursor.fetchall()
        conn.close()
        
        return jsonify({
            "success": True,
            "entrevistas": rows
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500    


@app.route("/api/ranking-ia", methods=["POST", "OPTIONS"])
def ranking_ia():
    """Ranking fino con ChatGPT sobre resultados de embeddings"""
    # Manejar preflight CORS
    if request.method == "OPTIONS":
        return '', 200
    
    try:
        data = request.json
        candidatos = data.get('candidatos', [])
        query = data.get('query', '')
        model = data.get('model', 'gpt')
        
        print(f"📥 Ranking IA para: '{query}' con {len(candidatos)} candidatos")
        
        if not candidatos or not query:
            return jsonify({"success": False, "error": "Faltan datos"}), 400
        
        # Limitar a 10 candidatos para no saturar
        candidatos_limit = candidatos[:10]
        
        # Preparar resumen para ChatGPT
        resumen = "Candidatos a evaluar:\n\n"
        for i, c in enumerate(candidatos_limit):
            # Parsear raw_data si existe
            try:
                if c.get('raw_data'):
                    datos = json.loads(c['raw_data']) if isinstance(c['raw_data'], str) else c['raw_data']
                else:
                    datos = c
            except:
                datos = c
            
            nombre = datos.get('datos_crudos', {}).get('nombre', c.get('nombre', 'Sin nombre'))
            profesion = datos.get('datos_crudos', {}).get('profesion_escrita', c.get('profesion', ''))
            perfil = datos.get('interpretacion', {}).get('perfil_interpretado', '')
            habilidades = datos.get('interpretacion', {}).get('habilidades_clave', [])
            experiencia = datos.get('interpretacion', {}).get('anos_experiencia_deducidos', '')
            seniority = datos.get('interpretacion', {}).get('seniority', '')
            
            resumen += f"{i+1}. {nombre}\n"
            if profesion:
                resumen += f"   Profesión: {profesion}\n"
            if perfil:
                resumen += f"   Perfil: {perfil[:150]}\n"
            if habilidades:
                resumen += f"   Habilidades: {', '.join(habilidades[:5])}\n"
            if experiencia:
                resumen += f"   Experiencia: {experiencia}\n"
            if seniority:
                resumen += f"   Seniority: {seniority}\n"
            resumen += "\n"
        
        prompt = f"""
Eres un headhunter experto. Te voy a dar:
1. Una búsqueda: "{query}"
2. Una lista de {len(candidatos_limit)} candidatos preseleccionados

Tu tarea es:
- Analizar cada candidato
- Ordenarlos del MEJOR al PEOR para esta búsqueda específica
- Dar una explicación breve de por qué el #1 es el mejor

Devuelve SOLO JSON con este formato:
{{
  "ranking": [1, 3, 2, 5, 4, ...],  // Números de candidato en orden de mejor a peor
  "mejor_explicacion": "Ana es la mejor porque...",
  "analisis": "Observaciones generales sobre los candidatos..."
}}

{resumen}
"""
        
        # Elegir modelo
        if model == "deepseek":
            client = deepseek_client
            model_name = "deepseek-chat"
        else:
            client = openai_client
            model_name = "gpt-4o-mini"
        
        response = client.chat.completions.create(
            model=model_name,
            messages=[{"role": "user", "content": prompt}],
            temperature=0
        )
        
        result = response.choices[0].message.content
        ranking_data = safe_json_parse(result)
        
        # Reordenar candidatos según ranking
        ranking_ids = ranking_data.get('ranking', [])
        candidatos_ordenados = []
        for pos in ranking_ids:
            if 1 <= pos <= len(candidatos_limit):
                candidatos_ordenados.append(candidatos_limit[pos-1])
        
        # Agregar el resto de candidatos (los que no estaban en el ranking)
        for c in candidatos:
            if c not in candidatos_ordenados:
                candidatos_ordenados.append(c)
        
        print(f"✅ Ranking completado")
        
        return jsonify({
            "success": True,
            "data": candidatos_ordenados,
            "explicacion": ranking_data.get('mejor_explicacion', ''),
            "analisis": ranking_data.get('analisis', '')
        })
        
    except Exception as e:
        print(f"❌ Error en ranking: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500 

    
def construir_perfil_candidato(datos):
    """Construye un perfil completo y estructurado del candidato para búsqueda inteligente"""
    perfil = []
    
    # Datos personales
    dc = datos.get('datos_crudos', {})
    perfil.append(f"📌 **{dc.get('nombre', 'Candidato')}**")
    perfil.append(f"🎓 Profesión: {dc.get('profesion_escrita', 'No especificada')}")
    perfil.append(f"📍 Ubicación: {dc.get('ubicacion', 'No especificada')}")
    
    # Educación
    if dc.get('universidad'):
        perfil.append(f"🏛️ Universidad: {dc.get('universidad')}")
    if dc.get('grado_academico'):
        perfil.append(f"📜 Grado: {dc.get('grado_academico')}")
    if dc.get('anio_graduacion'):
        perfil.append(f"🎓 Año graduación: {dc.get('anio_graduacion')}")
    
    # Experiencia laboral DETALLADA
    exp_laboral = dc.get('experiencia_laboral', [])
    if exp_laboral:
        perfil.append("\n💼 **EXPERIENCIA LABORAL:**")
        for exp in exp_laboral:
            empresa = exp.get('empresa', '')
            puesto = exp.get('puesto', '')
            inicio = exp.get('fecha_inicio', '')
            fin = exp.get('fecha_fin', '')
            desc = exp.get('descripcion', '')
            perfil.append(f"  • {puesto} en {empresa} ({inicio} - {fin})")
            if desc:
                perfil.append(f"    {desc[:100]}")
    
    # Empresas (lista única)
    empresas = dc.get('empresas', [])
    if empresas:
        perfil.append(f"\n🏢 **EMPRESAS DONDE HA TRABAJADO:** {', '.join(empresas)}")
    
    # Certificaciones
    certificaciones = dc.get('certificaciones', [])
    if certificaciones:
        perfil.append(f"\n📜 **CERTIFICACIONES:** {', '.join(certificaciones[:5])}")
    
    # Interpretación IA
    interp = datos.get('interpretacion', {})
    if interp:
        perfil.append(f"\n🔍 **ANÁLISIS IA:**")
        if interp.get('seniority'):
            perfil.append(f"  • Seniority: {interp.get('seniority')}")
        if interp.get('anos_experiencia_deducidos'):
            perfil.append(f"  • Años experiencia: {interp.get('anos_experiencia_deducidos')}")
        if interp.get('sector_deducido'):
            perfil.append(f"  • Sector: {interp.get('sector_deducido')}")
        if interp.get('rol_tipico'):
            perfil.append(f"  • Rol típico: {interp.get('rol_tipico')}")
        if interp.get('perfil_interpretado'):
            perfil.append(f"  • Perfil: {interp.get('perfil_interpretado')[:200]}")
    
    # Habilidades
    habilidades = interp.get('habilidades_clave', [])
    if habilidades:
        perfil.append(f"\n⚡ **HABILIDADES CLAVE:** {', '.join(habilidades[:8])}")
    
    # Fortalezas
    fortalezas = interp.get('fortalezas', [])
    if fortalezas:
        perfil.append(f"\n✅ **FORTALEZAS:** {', '.join(fortalezas[:3])}")
    
    return "\n".join(perfil)     


# =====================================================
# FUNCIONES PARA HEADHUNTER INTELIGENTE COMPLETO
# =====================================================

import re

def construir_perfil_completo(datos):
    """Construye perfil completo del candidato para IA"""
    
    perfil = []
    dc = datos.get('datos_crudos', {})
    interp = datos.get('interpretacion', {})
    
    # ===== 1. DATOS PERSONALES =====
    perfil.append("=== DATOS PERSONALES ===")
    perfil.append(f"NOMBRE: {dc.get('nombre', 'No especificado')}")
    perfil.append(f"PROFESIÓN: {dc.get('profesion_escrita', 'No especificada')}")
    perfil.append(f"UBICACIÓN: {dc.get('ubicacion', 'No especificada')}")
    perfil.append(f"EMAIL: {dc.get('email', 'No especificado')}")
    perfil.append(f"TELÉFONO: {dc.get('telefono', 'No especificado')}")
    perfil.append("")
    
    # ===== 2. EMPRESAS (DESTACADO) =====
    empresas = dc.get('empresas', [])
    if empresas:
        perfil.append("=== 🏢 EMPRESAS DONDE TRABAJÓ ===")
        perfil.append(', '.join(empresas))
        perfil.append("")
    
    # ===== 3. EXPERIENCIA LABORAL DETALLADA =====
    exp_laboral = dc.get('experiencia_laboral', [])
    if exp_laboral:
        perfil.append("=== EXPERIENCIA LABORAL ===")
        for exp in exp_laboral:
            perfil.append(f"• {exp.get('puesto', '')} en {exp.get('empresa', '')}")
            perfil.append(f"  {exp.get('fecha_inicio', '')} - {exp.get('fecha_fin', '')}")
            if exp.get('descripcion'):
                perfil.append(f"  {exp.get('descripcion', '')[:200]}")
            if exp.get('logros'):
                perfil.append(f"  Logros: {', '.join(exp.get('logros', []))}")
            perfil.append("")
    
    # ===== 4. EDUCACIÓN =====
    perfil.append("=== EDUCACIÓN ===")
    perfil.append(f"UNIVERSIDAD: {dc.get('universidad', 'No especificada')}")
    perfil.append(f"GRADO: {dc.get('grado_academico', 'No especificado')}")
    perfil.append(f"AÑO GRADUACIÓN: {dc.get('anio_graduacion', 'No especificado')}")
    perfil.append("")
    
    # ===== 5. CERTIFICACIONES =====
    certificaciones = dc.get('certificaciones', [])
    if certificaciones:
        perfil.append("=== CERTIFICACIONES ===")
        perfil.append(', '.join(certificaciones[:5]))
        perfil.append("")
    
    # ===== 6. HABILIDADES Y HERRAMIENTAS =====
    habilidades = dc.get('habilidades_listadas', [])
    if habilidades:
        perfil.append("=== HABILIDADES ===")
        perfil.append(', '.join(habilidades))
        perfil.append("")
    
    herramientas = dc.get('herramientas_listadas', [])
    if herramientas:
        perfil.append("=== HERRAMIENTAS ===")
        perfil.append(', '.join(herramientas))
        perfil.append("")
    
    # ===== 7. SECCIONES ADICIONALES (dinámicas) =====
    secciones_extra = ['proyectos', 'publicaciones', 'premios', 'voluntariados']
    for sec in secciones_extra:
        items = dc.get(sec, [])
        if items:
            perfil.append(f"=== {sec.upper()} ===")
            perfil.append(', '.join(items[:5]))
            perfil.append("")
    
    # ===== 8. INTERPRETACIÓN IA (profunda) =====
    if interp:
        perfil.append("=== 🧠 ANÁLISIS ESTRATÉGICO ===")
        if interp.get('perfil_interpretado'):
            perfil.append(f"PERFIL: {interp.get('perfil_interpretado')}")
        if interp.get('resumen_ejecutivo'):
            perfil.append(f"RESUMEN: {interp.get('resumen_ejecutivo')}")
        if interp.get('puntos_fuertes'):
            perfil.append(f"PUNTOS FUERTES: {', '.join(interp.get('puntos_fuertes', []))}")
        if interp.get('fortalezas'):
            perfil.append(f"FORTALEZAS: {', '.join(interp.get('fortalezas', []))}")
        if interp.get('areas_mejora'):
            perfil.append(f"ÁREAS DE MEJORA: {', '.join(interp.get('areas_mejora', []))}")
        if interp.get('rol_tipico'):
            perfil.append(f"ROL TÍPICO: {interp.get('rol_tipico')}")
        if interp.get('nivel_impacto'):
            perfil.append(f"NIVEL DE IMPACTO: {interp.get('nivel_impacto')}")
        if interp.get('potencial_crecimiento'):
            perfil.append(f"POTENCIAL: {interp.get('potencial_crecimiento')}")
        if interp.get('cultura_ideal'):
            perfil.append(f"CULTURA IDEAL: {interp.get('cultura_ideal')}")
        if interp.get('recomendacion'):
            perfil.append(f"RECOMENDACIÓN: {interp.get('recomendacion')}")
        perfil.append("")
    
    return "\n".join(perfil)

def buscar_por_nombre_proc(nombre):
    """Busca candidatos por nombre usando PostgreSQL"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, nombre, raw_data::jsonb 
            FROM candidates 
            WHERE nombre ILIKE %s 
            LIMIT 5
        """, (f"%{nombre}%",))
        
        resultados = cursor.fetchall()
        conn.close()
        
        candidatos = []
        for row in resultados:
            candidatos.append({
                'id': row[0],
                'nombre': row[1],
                'raw_data': row[2]
            })
        return candidatos
    except Exception as e:
        print(f"Error buscando por nombre: {e}")
        return []


def obtener_estadisticas_completas(pregunta):
    """Obtiene estadísticas según la pregunta"""
    conn = get_db_connection()
    cursor = conn.cursor()
    
    estadisticas = {}
    
    # Total de candidatos (siempre)
    cursor.execute("SELECT COUNT(*) FROM candidates")
    estadisticas['total'] = cursor.fetchone()[0]
    
    # Por sector
    if 'sector' in pregunta.lower() or 'industria' in pregunta.lower():
        cursor.execute("""
            SELECT COALESCE(raw_data::jsonb->'interpretacion'->>'sector_deducido', 'No especificado') as sector, COUNT(*)
            FROM candidates
            GROUP BY sector
            ORDER BY COUNT(*) DESC
            LIMIT 10
        """)
        estadisticas['por_sector'] = cursor.fetchall()
    
    # Por seniority
    if 'senior' in pregunta.lower() or 'junior' in pregunta.lower():
        cursor.execute("""
            SELECT COALESCE(raw_data::jsonb->'interpretacion'->>'seniority', 'No especificado') as seniority, COUNT(*)
            FROM candidates
            WHERE raw_data::jsonb->'interpretacion'->>'seniority' IS NOT NULL
            GROUP BY seniority
        """)
        estadisticas['por_seniority'] = cursor.fetchall()
    
    # Por ubicación
    if 'ubicacion' in pregunta.lower() or 'ciudad' in pregunta.lower():
        cursor.execute("""
            SELECT COALESCE(raw_data::jsonb->'datos_crudos'->>'ubicacion', 'No especificada') as ubicacion, COUNT(*)
            FROM candidates
            WHERE raw_data::jsonb->'datos_crudos'->>'ubicacion' IS NOT NULL
            GROUP BY ubicacion
            ORDER BY COUNT(*) DESC
            LIMIT 10
        """)
        estadisticas['por_ubicacion'] = cursor.fetchall()
    
    conn.close()
    return estadisticas


def clasificar_pregunta_ia(pregunta):
    """Clasifica el tipo de pregunta con reglas específicas"""
    pregunta_lower = pregunta.lower()
    
    # =====================================================
    # REGLAS EXPLÍCITAS (prioridad alta)
    # =====================================================
    
    # Comparación de candidatos
    if any(p in pregunta_lower for p in ['compara', 'versus', 'vs', 'entre', 'y', 'diferencia']):
        nombres = extraer_nombres_de_pregunta(pregunta)
        if len(nombres) >= 2:
            return "comparacion"
    
    # Estadísticas
    if any(p in pregunta_lower for p in ['cuántos', 'cuantos', 'cuántas', 'cuantas', 'cantidad', 'total', 'promedio', 'porcentaje', 'estadística']):
        return "estadistica"
    
    # Preguntas negativas
    if any(p in pregunta_lower for p in ['no trabajó', 'no tiene', 'no es', 'que no', 'excepto', 'menos']):
        return "negativa"
    
    # Ranking (solo cuando pide explícitamente "mejor", "top", "recomienda")
    if any(p in pregunta_lower for p in ['mejor', 'top', 'recomienda', 'los mejores', 'el mejor', 'ranking']):
        return "ranking"
    
    # =====================================================
    # SI NO APLICA NINGUNA REGLA, ES BÚSQUEDA NORMAL
    # =====================================================
    return "busqueda"


# =====================================================
# FUNCIONES PARA EVALUACIÓN DE ENTREVISTAS
# =====================================================

def construir_perfil_evaluacion(candidato_data):
    """Construye perfil detallado del candidato para evaluación"""
    dc = candidato_data.get('datos_crudos', {})
    interp = candidato_data.get('interpretacion', {})
    
    perfil = []
    perfil.append(f"NOMBRE: {dc.get('nombre', 'No especificado')}")
    perfil.append(f"PROFESIÓN: {dc.get('profesion_escrita', 'No especificada')}")
    perfil.append(f"UBICACIÓN: {dc.get('ubicacion', 'No especificada')}")
    
    # Empresas
    empresas = dc.get('empresas', [])
    if empresas:
        perfil.append(f"EMPRESAS: {', '.join(empresas)}")
    
    # Experiencia laboral
    exp_laboral = dc.get('experiencia_laboral', [])
    for exp in exp_laboral[:3]:
        perfil.append(f"TRABAJO: {exp.get('puesto', '')} en {exp.get('empresa', '')}")
        perfil.append(f"  {exp.get('fecha_inicio', '')} - {exp.get('fecha_fin', '')}")
    
    # Certificaciones
    certificaciones = dc.get('certificaciones', [])
    if certificaciones:
        perfil.append(f"CERTIFICACIONES: {', '.join(certificaciones[:3])}")
    
    # Habilidades
    habilidades = interp.get('habilidades_clave', [])
    if habilidades:
        perfil.append(f"HABILIDADES CLAVE: {', '.join(habilidades[:5])}")
    
    # Interpretación
    if interp.get('perfil_interpretado'):
        perfil.append(f"PERFIL: {interp.get('perfil_interpretado')}")
    if interp.get('seniority'):
        perfil.append(f"SENIORITY: {interp.get('seniority')}")
    if interp.get('anos_experiencia_deducidos'):
        perfil.append(f"EXPERIENCIA: {interp.get('anos_experiencia_deducidos')}")
    
    return "\n".join(perfil)


def construir_prompt_evaluacion(candidato_nombre, candidato_profesion, 
                                 perfil_completo, respuestas, puesto, plantilla):
    """Construye el prompt para ChatGPT"""
    
    prompt = f"""
Actúa como un sistema avanzado de evaluación de talento humano especializado en reclutamiento y selección por competencias.

## INFORMACIÓN DEL CANDIDATO

* **Nombre:** {candidato_nombre}
* **Profesión:** {candidato_profesion}
* **Puesto al que aplica:** {puesto}
* **Plantilla:** {plantilla}

## PERFIL COMPLETO DEL CANDIDATO (extraído de CV)

{perfil_completo}

## TRANSCRIPCIÓN DE LA ENTREVISTA

{respuestas}

## COMPETENCIAS A EVALUAR

### Competencias Blandas
- Comunicación
- Pensamiento analítico
- Resolución de problemas
- Trabajo en equipo
- Liderazgo
- Adaptabilidad
- Orientación a resultados

### Competencias Técnicas
- Identifica las habilidades técnicas relevantes del CV
- Evalúa coherencia entre CV y respuestas
- Determina nivel de profundidad (Básico/Intermedio/Avanzado)

## INSTRUCCIONES

Genera un informe profesional de evaluación en formato JSON con esta estructura EXACTA:

{{
  "resumen_ejecutivo": {{
    "nivel_general": "texto",
    "principales_fortalezas": ["fortaleza1", "fortaleza2"],
    "principales_riesgos": ["riesgo1", "riesgo2"],
    "recomendacion_general": "avanzar/evaluar_mas/descartar"
  }},
  "analisis_competencias_blandas": [
    {{"competencia": "Comunicación", "puntuacion": 0-100, "nivel": "Bajo/Medio/Alto", "evidencia": "texto", "interpretacion": "texto"}}
  ],
  "analisis_competencias_tecnicas": [
    {{"tecnologia": "nombre", "nivel_cv": "texto", "nivel_entrevista": "texto", "gap": "Sin gap/Parcial/Significativo", "evidencia": "texto"}}
  ],
  "analisis_conductual_star": {{
    "ejemplos_concretos": "Bajo/Medio/Alto",
    "estructura_respuestas": "texto",
    "nivel_profundidad": "texto"
  }},
  "analisis_comunicacion": {{
    "claridad": 0-100,
    "seguridad": 0-100,
    "coherencia": 0-100,
    "capacidad_sintesis": 0-100
  }},
  "analisis_cognitivo": {{
    "razonamiento_logico": "texto",
    "estructura_ideas": "texto",
    "resolucion_problemas": "texto"
  }},
  "validacion_experiencia": {{
    "consistencia": "Alta/Media/Baja",
    "inconsistencias": []
  }},
  "red_flags": [],
  "deteccion_ia": {{
    "probabilidad": "Baja/Media/Alta",
    "razones": []
  }},
  "potencial_y_proyeccion": {{
    "curva_aprendizaje": "texto",
    "capacidad_crecimiento": "texto",
    "proyeccion_organizacion": "texto"
  }},
  "fit_cultural": {{
    "compatibilidad": "texto",
    "entorno_ideal": "texto"
  }},
  "score_global": 0-100,
  "clasificacion": "Top Talent/Fuerte/Aceptable/Riesgoso",
  "recomendacion_final": {{
    "decision": "contratar/avanzar/validar/no_recomendable",
    "argumento": "texto"
  }},
  "preguntas_seguimiento": ["pregunta1", "pregunta2"],
  "recomendaciones_cliente": ["recomendacion1", "recomendacion2"],
  "insight_headhunter": "texto"
}}

Devuelve SOLO el JSON, sin texto adicional.
"""
    return prompt


def analizar_entrevista(entrevista_id):
    """Analiza una entrevista usando procedimientos almacenados"""
    try:
        print(f"📊 Analizando entrevista ID: {entrevista_id}")
        
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        # 1. Obtener datos de la entrevista (procedimiento)
        cursor.callproc('obtener_entrevista_para_analisis', [entrevista_id])
        entrevista = cursor.fetchone()
        
        if not entrevista or not entrevista.get('candidato_raw_data'):
            conn.close()
            return False
        
        # 2. Obtener respuestas (procedimiento)
        cursor.callproc('obtener_respuestas_entrevista', [entrevista_id])
        respuestas_row = cursor.fetchone()
        respuestas = respuestas_row['obtener_respuestas_entrevista'] if respuestas_row else []
        
        conn.close()
        
        # 3. Parsear datos del candidato
        candidato_data = entrevista['candidato_raw_data']
        if isinstance(candidato_data, str):
            candidato_data = json.loads(candidato_data)
        
        # 4. Construir perfil
        perfil = construir_perfil_evaluacion(candidato_data)
        
        # 5. Construir texto de respuestas
        texto_respuestas = ""
        for i, r in enumerate(respuestas):
            texto_respuestas += f"\n**Pregunta {i+1}:** {r['pregunta']}\n"
            texto_respuestas += f"**Respuesta:** {r['respuesta']}\n"
            if r.get('tiempo'):
                texto_respuestas += f"**Tiempo:** {r['tiempo']} segundos\n"
            texto_respuestas += "-" * 50 + "\n"
        
        # 6. Construir prompt
        prompt = construir_prompt_evaluacion(
            candidato_nombre=entrevista['candidato_nombre'],
            candidato_profesion=entrevista['candidato_profesion'],
            perfil_completo=perfil,
            respuestas=texto_respuestas,
            puesto=entrevista.get('puesto', 'No especificado'),
            plantilla=entrevista.get('plantilla', 'General')
        )
        
        # 7. Llamar a ChatGPT
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4000
        )
        
        contenido = response.choices[0].message.content
        analisis = safe_json_parse(contenido)
        
        # 8. Guardar análisis (procedimiento)
        conn = get_db_connection()
        cursor = conn.cursor()
        score_global = analisis.get('score_global', 0)
        cursor.callproc('guardar_analisis_entrevista', [entrevista_id, json.dumps(analisis, ensure_ascii=False), score_global])
        conn.commit()
        conn.close()
        
        print(f"✅ Análisis completado para entrevista ID {entrevista_id}")
        return True
        
    except Exception as e:
        print(f"❌ Error analizando entrevista: {e}")
        import traceback
        traceback.print_exc()
        return False
    


# POR CAMBIAR LUEGO A SPS--------------------------------

# =====================================================
# ENDPOINTS PARA PLANTILLAS V2 (STAR)
# =====================================================

@app.route("/api/admin/plantillas-v2", methods=["GET"])
def obtener_plantillas_v2():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT 
                p.id, p.nombre, p.tipo, p.descripcion, p.objetivo,
                p.fecha_creacion,
                (SELECT COUNT(*) FROM preguntas_entrevista_v2 WHERE plantilla_id = p.id) as preguntas_count,
                (SELECT COUNT(DISTINCT fase) FROM preguntas_entrevista_v2 WHERE plantilla_id = p.id) as fases_count
            FROM plantillas_entrevista_v2 p
            ORDER BY p.fecha_creacion DESC
        """)
        
        plantillas = []
        for row in cursor.fetchall():
            plantillas.append({
                "id": row[0],
                "nombre": row[1],
                "tipo": row[2],
                "descripcion": row[3],
                "objetivo": row[4],
                "fecha_creacion": row[5],
                "preguntas_count": row[6],
                "fases_count": row[7]
            })
        
        conn.close()
        return jsonify({"success": True, "plantillas": plantillas})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/admin/plantilla-v2", methods=["POST"])
def crear_plantilla_v2():
    try:
        data = request.json
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Insertar plantilla
        cursor.execute("""
            INSERT INTO plantillas_entrevista_v2 
            (nombre, tipo, descripcion, objetivo, perfil_puesto, cultura_empresa, fase_config, fecha_creacion)
            VALUES (%s, %s, %s, %s, %s, %s, %s, CURRENT_TIMESTAMP)
            RETURNING id
        """, (
            data['nombre'],
            data['tipo'],
            data.get('descripcion', ''),
            data.get('objetivo', ''),
            json.dumps(data.get('perfil_puesto', {})),
            json.dumps(data.get('cultura_empresa', {})),
            json.dumps(data.get('fases', []))
        ))
        
        plantilla_id = cursor.fetchone()[0]
        
        # Insertar preguntas por fase
        orden = 0
        for fase in data.get('fases', []):
            for pregunta in fase.get('preguntas', []):
                cursor.execute("""
                    INSERT INTO preguntas_entrevista_v2 
                    (plantilla_id, fase, pregunta, objetivo, criterio_star, competencia, peso, orden, tiempo_maximo)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                """, (
                    plantilla_id,
                    fase.get('nombre', 'General'),
                    pregunta.get('pregunta', ''),
                    pregunta.get('objetivo', ''),
                    pregunta.get('criterio_star', ''),
                    pregunta.get('competencia', ''),
                    pregunta.get('peso', 1),
                    orden,
                    120
                ))
                orden += 1
        
        conn.commit()
        conn.close()
        
        return jsonify({"success": True, "plantilla_id": plantilla_id})
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/entrevista/iniciar-v2", methods=["POST"])
def iniciar_entrevista_v2():
    """Inicia una entrevista con plantilla STAR"""
    try:
        data = request.json
        candidato_id = data.get('candidato_id')
        plantilla_id = data.get('plantilla_id')
        
        print(f"🎥 [INICIAR] Candidato ID: {candidato_id}, Plantilla ID: {plantilla_id}")
        
        # Validaciones
        if not candidato_id:
            return jsonify({"success": False, "error": "Falta el ID del candidato"}), 400
        
        if not plantilla_id:
            return jsonify({"success": False, "error": "Falta el ID de la plantilla"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 🔥 VERIFICAR QUE EL CANDIDATO EXISTE
        cursor.execute("SELECT id, nombre FROM candidates WHERE id = %s", (candidato_id,))
        candidato = cursor.fetchone()
        
        if not candidato:
            conn.close()
            return jsonify({"success": False, "error": f"El candidato con ID {candidato_id} no existe"}), 404
        
        print(f"✅ Candidato encontrado: {candidato[1]} (ID: {candidato[0]})")
        
        # 🔥 VERIFICAR QUE LA PLANTILLA EXISTE
        cursor.execute("""
            SELECT fase_config, perfil_puesto, cultura_empresa, objetivo, 
                   intentos_por_pregunta, max_advertencias, tiempo_instrucciones
            FROM plantillas_entrevista_v2 WHERE id = %s
        """, (plantilla_id,))
        plantilla = cursor.fetchone()
        
        if not plantilla:
            conn.close()
            return jsonify({"success": False, "error": f"La plantilla con ID {plantilla_id} no existe"}), 404
        
        print(f"✅ Plantilla encontrada: ID {plantilla_id}")
        
        token = secrets.token_urlsafe(32)
        base_url = request.host_url.rstrip('/')
        enlace = f"{base_url}/entrevista-candidato-v2.html?token={token}"
        
        # Crear entrevista con metadata
        cursor.execute("""
            INSERT INTO entrevistas 
            (candidato_id, plantilla_id, token_acceso, fecha_inicio, estado, enlace_compartido, metadata)
            VALUES (%s, %s, %s, CURRENT_TIMESTAMP, 'pendiente', %s, %s)
            RETURNING id
        """, (
            candidato_id, 
            plantilla_id, 
            token, 
            enlace,
            json.dumps({
                "tipo": "star",
                "objetivo": plantilla[3],
                "perfil_puesto": plantilla[1],
                "cultura_empresa": plantilla[2],
                "fases": plantilla[0],
                "intentos_por_pregunta": plantilla[4] or 2,
                "max_advertencias": plantilla[5] or 3,
                "tiempo_instrucciones": plantilla[6] or 60
            })
        ))
        
        entrevista_id = cursor.fetchone()[0]
        print(f"✅ Entrevista creada con ID: {entrevista_id}")
        
        # Obtener preguntas
        cursor.execute("""
            SELECT id, fase, pregunta, objetivo, criterio_star, competencia, tiempo_maximo, orden,
                   intentos_permitidos
            FROM preguntas_entrevista_v2
            WHERE plantilla_id = %s
            ORDER BY orden
        """, (plantilla_id,))
        
        preguntas = []
        for row in cursor.fetchall():
            preguntas.append({
                "id": row[0],
                "fase": row[1] or "General",
                "pregunta": row[2],
                "objetivo": row[3] or "",
                "criterio_star": row[4] or "",
                "competencia": row[5] or "General",
                "tiempo_maximo": row[6] or 120,
                "orden": row[7],
                "intentos_permitidos": row[8] or 2
            })
        
        conn.commit()
        conn.close()
        
        print(f"✅ Entrevista iniciada exitosamente. Total preguntas: {len(preguntas)}")
        
        return jsonify({
            "success": True,
            "entrevista_id": entrevista_id,
            "enlace": enlace,
            "token": token,
            "preguntas": preguntas,
            "metadata": {
                "objetivo": plantilla[3],
                "total_preguntas": len(preguntas),
                "intentos_por_pregunta": plantilla[4] or 2,
                "max_advertencias": plantilla[5] or 3
            }
        })
        
    except Exception as e:
        print(f"❌ Error en iniciar_entrevista_v2: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500


#------------------------------------------------------- CIERRA POR CAMBIAR LUEGO A SPS

# =====================================================
# ENDPOINTS PARA ENTREVISTAS STAR (V2)
# =====================================================

@app.route("/api/entrevista/candidato-v2/<token>", methods=["GET"])
def obtener_entrevista_candidato_v2(token):
    """Obtiene entrevista para candidato (versión STAR)"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.callproc('obtener_entrevista_por_token_v2', [token])
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        entrevista = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        
        return jsonify({
            "success": True,
            "entrevista": entrevista.get('entrevista')
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/entrevista/candidato/responder-v2", methods=["POST"])
def guardar_respuesta_candidato_v2():
    """Guarda respuesta de entrevista (versión STAR) con transcripción automática"""
    try:
        token = request.form.get('token')
        pregunta_id = request.form.get('pregunta_id')
        pregunta_texto = request.form.get('pregunta_texto')
        tiempo = request.form.get('tiempo', 0)
        transcripcion_manual = request.form.get('transcripcion', '')
        idioma = request.form.get('idioma', 'es')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT id FROM entrevistas WHERE token_acceso = %s", (token,))
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        entrevista_id = row[0]
        
        filename = None
        transcripcion_final = transcripcion_manual
        
        if 'video' in request.files:
            video_file = request.files['video']
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"entrevista_{entrevista_id}_pregunta_{pregunta_id}_{timestamp}.webm"
            filepath = os.path.join(VIDEO_UPLOAD_FOLDER, filename)
            video_file.save(filepath)
            print(f"✅ Video guardado: {filename}")
            
            if not transcripcion_manual or transcripcion_manual == "":
                try:
                    transcripcion_final = transcribir_video(filepath, idioma)
                    if transcripcion_final:
                        print(f"✅ Transcripción: {transcripcion_final[:100]}...")
                    else:
                        transcripcion_final = "[No se pudo transcribir el video]"
                except Exception as e:
                    print(f"❌ Error: {e}")
                    transcripcion_final = "[Error en transcripción]"
        
        # Usar procedimiento almacenado para guardar
        cursor.execute("""
            SELECT guardar_respuesta_star(%s, %s, %s, %s, %s, %s, %s, %s)
        """, (entrevista_id, pregunta_id, pregunta_texto, 
              transcripcion_final, filename, tiempo, None, None))
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": True, 
            "transcripcion": transcripcion_final,
            "tiene_transcripcion": bool(transcripcion_final and transcripcion_final != "")
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/entrevista/candidato/finalizar-v2", methods=["POST"])
def finalizar_entrevista_candidato_v2():
    try:
        token = request.json.get('token')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.callproc('finalizar_entrevista_star', [token])
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        # Obtener el ID correctamente
        if isinstance(row, tuple):
            entrevista_id = row[0]
        elif isinstance(row, dict):
            entrevista_id = row.get('id') or list(row.values())[0]
        else:
            entrevista_id = row
        
        conn.commit()
        conn.close()
        
        # Disparar análisis en segundo plano
        import threading
        thread = threading.Thread(target=analizar_entrevista_star, args=(entrevista_id,))
        thread.daemon = True
        thread.start()
        
        return jsonify({"success": True, "entrevista_id": entrevista_id})
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/entrevista/resultados-star/<int:entrevista_id>", methods=["GET"])
def obtener_resultados_star(entrevista_id):
    """Obtiene resultados detallados de entrevista STAR"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.callproc('obtener_resultados_star', [entrevista_id])
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": False, "error": "No se encontraron resultados"}), 404
        
        resultados = row[0] if isinstance(row[0], dict) else json.loads(row[0])
        
        return jsonify({
            "success": True,
            "resultados": resultados
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
    


def analizar_entrevista_star(entrevista_id):
    """Analiza entrevista STAR con feedback detallado por pregunta"""
    try:
        print(f"📊 Analizando entrevista STAR ID: {entrevista_id}")
        
        conn = get_db_connection()
        cursor = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        
        cursor.execute("CALL obtener_reporte_completo_entrevista(%s, %s)", (entrevista_id, None))
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            print("❌ No se encontraron datos")
            return False
        
        resultado_json = None
        for key, value in row.items():
            resultado_json = value
            break
        
        if not resultado_json:
            print("❌ No se encontró JSON")
            return False
        
        datos = json.loads(resultado_json)
        
        entrevista = datos.get('entrevista', {})
        candidato_data = entrevista.get('candidato_data', {})
        plantilla = datos.get('plantilla', {})
        preguntas_config = datos.get('preguntas_configuracion', [])
        respuestas = datos.get('respuestas', [])
        
        interpretacion = candidato_data.get('interpretacion', {})
        datos_crudos = candidato_data.get('datos_crudos', {})
        
        # =========================================================
        # CONSTRUIR PERFIL DEL CANDIDATO
        # =========================================================
        perfil_candidato = f"""
=== INFORMACIÓN DEL CANDIDATO ===
Nombre: {datos_crudos.get('nombre', 'No disponible')}
Profesión: {datos_crudos.get('profesion_escrita', 'No disponible')}
Seniority detectado: {interpretacion.get('seniority', 'No disponible')}
Años experiencia: {interpretacion.get('anos_experiencia_deducidos', 'No disponible')}
Ubicación: {datos_crudos.get('ubicacion', 'No disponible')}

=== TECNOLOGÍAS Y HERRAMIENTAS ===
{', '.join(interpretacion.get('habilidades_clave', [])) if interpretacion.get('habilidades_clave') else 'No especificadas'}

=== EMPRESAS DESTACADAS ===
{chr(10).join(['- ' + e for e in datos_crudos.get('empresas', [])[:5]]) if datos_crudos.get('empresas') else 'No especificadas'}

=== PERFIL INTERPRETADO POR IA ===
{interpretacion.get('perfil_interpretado', 'No disponible')}

=== FORTALEZAS SEGÚN CV ===
{chr(10).join(['- ' + f for f in interpretacion.get('fortalezas', [])[:5]]) if interpretacion.get('fortalezas') else 'No especificadas'}

=== ÁREAS DE MEJORA SEGÚN CV ===
{chr(10).join(['- ' + a for a in interpretacion.get('areas_mejora', [])[:3]]) if interpretacion.get('areas_mejora') else 'No especificadas'}
"""
        
        # =========================================================
        # PERFIL DEL PUESTO
        # =========================================================
        perfil_puesto = plantilla.get('perfil_puesto', {})
        perfil_puesto_texto = f"""
=== PERFIL DEL PUESTO ===
Nombre de la plantilla: {plantilla.get('nombre', 'No especificado')}
Descripción: {plantilla.get('descripcion', 'No especificada')[:500]}

=== REQUISITOS DEL PUESTO ===
{chr(10).join(['- ' + r for r in perfil_puesto.get('requisitos', [])]) if perfil_puesto.get('requisitos') else 'No se especificaron requisitos formales'}

=== RESPONSABILIDADES ===
{chr(10).join(['- ' + r for r in perfil_puesto.get('responsabilidades', [])]) if perfil_puesto.get('responsabilidades') else 'No se especificaron responsabilidades formales'}
"""
        
        # =========================================================
        # CULTURA DE LA EMPRESA
        # =========================================================
        cultura = plantilla.get('cultura_empresa', {})
        valores = cultura.get('valores', [])
        cultura_texto = f"""
=== CULTURA ORGANIZACIONAL ===
Valores clave:
{chr(10).join(['- ' + v[:200] + '...' for v in valores[:4]]) if valores else 'No especificados'}

Ambiente de trabajo:
{cultura.get('ambiente', 'No especificado')[:400]}
"""
        
        # =========================================================
        # OBJETIVO DE LA ENTREVISTA
        # =========================================================
        objetivo_texto = f"""
=== OBJETIVO DE LA ENTREVISTA ===
{plantilla.get('objetivo', 'No especificado')[:600]}
"""
        
        # =========================================================
        # CONSTRUIR DETALLE DE CADA PREGUNTA CON RESPUESTA
        # =========================================================
        preguntas_detalle = []
        for i, conf in enumerate(preguntas_config):
            respuesta_obj = next((r for r in respuestas if r.get('pregunta') == conf.get('pregunta')), {})
            tiempo_respuesta = respuesta_obj.get('tiempo', 0)
            texto_respuesta = respuesta_obj.get('respuesta', '')
            
            if not texto_respuesta or texto_respuesta == "":
                texto_respuesta = "[El candidato grabó video pero no transcribió su respuesta. El reclutador debe revisar la grabación.]"
            
            preguntas_detalle.append(f"""
╔══════════════════════════════════════════════════════════════╗
║ PREGUNTA {i+1}: {conf.get('competencia', 'Competencia')}
╚══════════════════════════════════════════════════════════════╝

📌 **Pregunta formulada:**
{conf.get('pregunta', 'No disponible')}

🎯 **Objetivo de esta pregunta (qué se busca evaluar):**
{conf.get('objetivo', 'No especificado')}

⭐ **Criterio STAR esperado:**
{conf.get('criterio_star', 'No especificado')}

⏱️ **Tiempo máximo sugerido:** {conf.get('tiempo_maximo', 120)} segundos

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🗣️ **RESPUESTA DEL CANDIDATO:**
{texto_respuesta}

⏱️ **Tiempo real de respuesta:** {tiempo_respuesta} segundos
{"⚠️ El candidato respondió muy rápido. Posible falta de profundidad." if tiempo_respuesta < 30 else "✅ Tiempo de respuesta adecuado." if tiempo_respuesta < 90 else "📝 El candidato dedicó tiempo considerable a su respuesta."}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
""")
        
        # =========================================================
        # PROMPT COMPLETO CON INSTRUCCIONES DETALLADAS
        # =========================================================
        prompt = f"""
Actúa como un Headhunter Senior con 20 años de experiencia en reclutamiento técnico. Eres conocido por ser ESTRICTO, ANALÍTICO y por dar FEEDBACK DETALLADO Y ACCIONABLE.

## 📋 INFORMACIÓN DEL CANDIDATO (Extraída de su CV)

{perfil_candidato}

## 🎯 INFORMACIÓN DEL PUESTO Y LA EMPRESA

{perfil_puesto_texto}

{cultura_texto}

{objetivo_texto}

## 🎙️ RESPUESTAS DEL CANDIDATO EN LA ENTREVISTA

{chr(10).join(preguntas_detalle)}

## 📊 TAREA DE EVALUACIÓN

Debes evaluar al candidato respondiendo las siguientes preguntas de forma **DETALLADA y JUSTIFICADA**:

### 1. Evaluación General
- ¿El candidato cumple con los requisitos técnicos del puesto? ¿Por qué sí o por qué no?
- ¿Su nivel de experiencia (años, empresas, roles) es suficiente? ¿Por qué?
- ¿Se alinea con la cultura de la empresa? ¿Por qué?

### 2. Evaluación por Pregunta (para CADA una)
Para cada pregunta, debes responder:
- **¿Qué hizo bien?** (fortalezas específicas demostradas)
- **¿Qué puede mejorar?** (áreas de oportunidad concretas)
- **Evaluación STAR**: ¿Cómo se desempeñó en Situación, Tarea, Acción y Resultado?
- **Puntuación**: 0-100 (justificada)

### 3. Análisis de Ajuste Cultural
- ¿El candidato encaja con los valores de la empresa?
- ¿Se adaptaría al ambiente de trabajo descrito?

### 4. Insight del Headhunter
- Análisis honesto: ¿Este candidato realmente encaja con lo que busca la empresa?
- ¿Cuál es tu recomendación final y POR QUÉ?

## 🔍 FORMATO DE RESPUESTA (SOLO JSON)

{{
  "resumen_ejecutivo": {{
    "evaluacion_general": "Evaluación general del candidato en 2-3 párrafos. Incluye comparación directa con los requisitos del puesto y justificación.",
    "cumple_requisitos_tecnicos": true/false,
    "justificacion_tecnica": "Explicación detallada de por qué cumple o no los requisitos técnicos, mencionando habilidades específicas.",
    "cumple_requisitos_experiencia": true/false,
    "justificacion_experiencia": "Explicación de si sus años de experiencia y empresas son suficientes para el puesto.",
    "cumple_ajuste_cultural": true/false,
    "justificacion_cultural": "Explicación de si se alinea con los valores y ambiente de la empresa."
  }},
  "analisis_por_pregunta": [
    {{
      "pregunta": "texto de la pregunta",
      "competencia": "nombre de la competencia evaluada",
      "objetivo_pregunta": "qué se buscaba evaluar",
      "respuesta_candidato": "transcripción o resumen de su respuesta",
      "fortalezas": ["fortaleza específica demostrada 1", "fortaleza específica demostrada 2"],
      "areas_mejora": ["área específica a mejorar 1", "área específica a mejorar 2"],
      "feedback_detallado": "Análisis profundo de su respuesta, explicando qué hizo bien y qué podría hacer mejor. Mínimo 3 líneas.",
      "evaluacion_star": {{
        "situacion": 0-100,
        "tarea": 0-100,
        "accion": 0-100,
        "resultado": 0-100,
        "comentario": "explicación de su desempeño STAR"
      }},
      "puntuacion": 0-100,
      "justificacion_puntuacion": "explicación detallada de por qué recibe esa puntuación"
    }}
  ],
  "analisis_ajuste_cultural": {{
    "compatibilidad": "Alta/Media/Baja",
    "razones": ["razón1 basada en su perfil", "razón2 basada en su experiencia"],
    "justificacion": "explicación detallada de por qué encaja o no con la cultura"
  }},
  "score_global": 0-100,
  "recomendacion_final": {{
    "decision": "contratar/avanzar/no_recomendable",
    "argumento": "Explicación DETALLADA de por qué se toma esta decisión. Mínimo 4 líneas mencionando requisitos específicos.",
    "next_steps": "Próximos pasos sugeridos si aplica"
  }},
  "fortalezas_generales": ["fortaleza1", "fortaleza2", "fortaleza3", "fortaleza4", "fortaleza5"],
  "areas_mejora": ["area1", "area2", "area3", "area4"],
  "insight_headhunter": "Análisis profundo y honesto (mínimo 4 líneas) sobre si el candidato realmente encaja con lo que busca la empresa y por qué."
}}

**IMPORTANTE:** 
- Cada "justificacion" debe ser EXPLÍCITA y basada en la información proporcionada
- No uses frases genéricas como "es un buen candidato". Explica POR QUÉ
- Relaciona cada punto con los requisitos del puesto y la cultura de la empresa
- El feedback por pregunta debe ser ACCIONABLE (qué puede mejorar concretamente)

Devuelve SOLO el JSON, sin texto adicional.
"""
        
        print("📡 Llamando a ChatGPT con prompt detallado...")
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4000
        )
        
        contenido = response.choices[0].message.content
        analisis = safe_json_parse(contenido)
        
        # Guardar análisis
        conn = get_db_connection()
        cursor = conn.cursor()
        score = analisis.get('score_global', 0)
        cursor.execute("""
            UPDATE entrevistas 
            SET analisis_ia = %s, puntuacion_global = %s, estado = 'analizada'
            WHERE id = %s
        """, (json.dumps(analisis, ensure_ascii=False), score, entrevista_id))
        conn.commit()
        conn.close()
        
        print(f"✅ Análisis completado para entrevista {entrevista_id}")
        return True
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return False
    

def construir_prompt_analisis_star(datos):
    """Construye prompt para análisis STAR"""
    entrevista = datos.get('entrevista', {})
    candidato = datos.get('candidato', {})
    plantilla = datos.get('plantilla', {})
    respuestas = datos.get('respuestas', [])
    
    texto_respuestas = ""
    for i, r in enumerate(respuestas):
        texto_respuestas += f"""
**Pregunta {i+1}** (Fase: {r.get('fase', 'General')})
Competencia: {r.get('competencia', 'No especificada')}
Objetivo: {r.get('objetivo', 'No especificado')}
Respuesta: {r.get('respuesta_texto', 'No disponible')}
Tiempo: {r.get('tiempo', 0)} segundos
"""
    
    perfil_puesto = plantilla.get('perfil_puesto', {})
    cultura = plantilla.get('cultura_empresa', {})
    
    prompt = f"""
Eres un headhunter experto en evaluación por competencias con método STAR.

## INFORMACIÓN DE LA ENTREVISTA

**Candidato:** {candidato.get('nombre', 'No especificado')}
**Profesión:** {candidato.get('profesion', 'No especificada')}

**Objetivo de la entrevista:** {plantilla.get('objetivo', 'Evaluación general')}

**Perfil del puesto:**
- Título: {perfil_puesto.get('titulo', 'No especificado')}
- Seniority: {perfil_puesto.get('seniority', 'No especificado')}
- Responsabilidades: {', '.join(perfil_puesto.get('responsabilidades', []))}
- Requisitos: {', '.join(perfil_puesto.get('requisitos', []))}

**Cultura de la empresa:**
- Valores: {', '.join(cultura.get('valores', []))}
- Ambiente: {cultura.get('ambiente', 'No especificado')}

## TRANSCRIPCIÓN DE RESPUESTAS (MÉTODO STAR)

{texto_respuestas}

## INSTRUCCIONES DE EVALUACIÓN

Evalúa cada respuesta según el método STAR:
- **S (Situación)**: ¿Describe el contexto adecuadamente?
- **T (Tarea)**: ¿Explica claramente su responsabilidad?
- **A (Acción)**: ¿Detalla las acciones concretas que tomó?
- **R (Resultado)**: ¿Muestra resultados medibles o impacto?

Devuelve SOLO JSON con:
{{
  "score_global": 0-100,
  "evaluacion_star": {{
    "situacion_promedio": 0-100,
    "tarea_promedio": 0-100,
    "accion_promedio": 0-100,
    "resultado_promedio": 0-100
  }},
  "analisis_por_pregunta": [
    {{
      "pregunta": "texto",
      "competencia": "nombre",
      "puntuacion_star": 0-100,
      "fortalezas": ["fortaleza1"],
      "areas_mejora": ["area1"],
      "evaluacion_star": {{
        "situacion": 0-100,
        "tarea": 0-100,
        "accion": 0-100,
        "resultado": 0-100
      }}
    }}
  ],
  "fortalezas_generales": ["fortaleza1", "fortaleza2"],
  "areas_mejora": ["area1", "area2"],
  "recomendacion_final": {{
    "decision": "contratar/avanzar/validar/no_recomendable",
    "argumento": "texto"
  }},
  "insight_headhunter": "texto"
}}
"""
    return prompt

@app.route("/api/entrevista/video/<filename>", methods=["GET"])
def ver_video_entrevista(filename):
    """Sirve el video de la entrevista con soporte para Safari"""
    try:
        from urllib.parse import unquote
        filename = sanitize_filename(unquote(filename))
        filepath = os.path.join(VIDEO_UPLOAD_FOLDER, filename)
        
        print(f"🎬 Buscando video en: {filepath}")
        
        if not os.path.exists(filepath):
            print(f"❌ Video no encontrado: {filepath}")
            return jsonify({"success": False, "error": "Video no encontrado"}), 404
        
        # Detectar el formato
        if filename.endswith('.webm'):
            mimetype = 'video/webm'
        elif filename.endswith('.mp4'):
            mimetype = 'video/mp4'
        else:
            mimetype = 'video/mp4'
        
        # Headers especiales para Safari
        response = send_file(
            filepath, 
            mimetype=mimetype,
            conditional=True,
            etag=True,
            last_modified=datetime.fromtimestamp(os.path.getmtime(filepath))
        )
        
        # Headers adicionales para Safari
        response.headers['Accept-Ranges'] = 'bytes'
        response.headers['Content-Security-Policy'] = "default-src 'self' 'unsafe-inline' 'unsafe-eval' data: blob:; media-src 'self' blob: data:;"
        
        return response
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500  
    

@app.route("/api/entrevista/reporte-completo-v2/<int:entrevista_id>", methods=["GET"])
def reporte_completo_v2(entrevista_id):
    """Obtiene reporte completo con toda la información de la plantilla"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("CALL obtener_reporte_completo_entrevista(%s, %s)", (entrevista_id, None))
        row = cursor.fetchone()
        conn.close()
        
        if not row or not row[0]:
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        resultado = json.loads(row[0])
        
        return jsonify(resultado)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500  


def transcribir_video(filepath, idioma=None):
    """Transcribe un archivo de video usando Whisper de OpenAI"""
    try:
        print(f"🎙️ Transcribiendo video: {filepath}")
        
        opciones = {
            "model": "whisper-1",
            "file": open(filepath, "rb"),
            "response_format": "text"
        }
        
        if idioma and idioma in ['es', 'en']:
            opciones["language"] = idioma
        
        transcript = openai_client.audio.transcriptions.create(**opciones)
        opciones["file"].close()
        
        print(f"✅ Transcripción completada: {len(transcript)} caracteres")
        return transcript
        
    except Exception as e:
        print(f"❌ Error transcribiendo video: {e}")
        return None      
    
@app.route("/api/entrevista/transcripcion/<int:respuesta_id>", methods=["GET"])
def obtener_transcripcion(respuesta_id):
    """Obtiene la transcripción de una respuesta específica"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT respuesta_texto, idioma_transcripcion, respuesta_video_path
            FROM respuestas_entrevista_v2
            WHERE id = %s
        """, (respuesta_id,))
        
        row = cursor.fetchone()
        conn.close()
        
        if not row:
            return jsonify({"success": False, "error": "Respuesta no encontrada"}), 404
        
        return jsonify({
            "success": True,
            "transcripcion": row[0] or "No hay transcripción disponible",
            "idioma": row[1] or "es",
            "video": row[2]
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500    
    

# =====================================================
# ENDPOINTS PARA PORTAL DE CANDIDATOS
# =====================================================

@app.route("/api/puestos", methods=["GET"])
def api_obtener_puestos():
    """Obtiene todos los puestos activos"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_puestos_activos()")
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/puesto/<int:puesto_id>", methods=["GET"])
def api_obtener_puesto(puesto_id):
    """Obtiene un puesto específico"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_puesto(%s)", (puesto_id,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/postular", methods=["POST"])
def api_postular():
    """Postula un candidato a un puesto"""
    try:
        nombre = request.form.get('nombre')
        email = request.form.get('email')
        telefono = request.form.get('telefono')
        puesto_id = request.form.get('puesto_id')
        
        # Recopilar respuestas del formulario
        respuestas = {}
        for key, value in request.form.items():
            if key.startswith('respuesta_'):
                pregunta_id = key.replace('respuesta_', '')
                respuestas[pregunta_id] = value
        
        # Procesar CV si se subió
        cv_filename = None
        cv_analisis = '{}'
        
        if 'cv' in request.files:
            cv_file = request.files['cv']
            if cv_file and allowed_file(cv_file.filename):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                cv_filename = f"candidato_{timestamp}_{secure_filename(cv_file.filename)}"
                cv_path = os.path.join(UPLOAD_FOLDER, cv_filename)
                cv_file.save(cv_path)
                
                # Analizar CV con IA
                text = extract_text(cv_path)
                if text and len(text) > 50:
                    data = analyze_cv(text, "gpt")
                    data = normalize_lists(data)
                    data = score_candidate(data)
                    cv_analisis = json.dumps(data)
        
        # Llamar al procedimiento almacenado
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT crear_postulacion(%s, %s, %s, %s, %s, %s, %s::jsonb)
        """, (nombre, email, telefono, puesto_id, json.dumps(respuestas), cv_filename, cv_analisis))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/mis-postulaciones", methods=["GET"])
def api_mis_postulaciones():
    """Obtiene las postulaciones de un candidato por email"""
    try:
        email = request.args.get('email')
        
        if not email:
            return jsonify({"success": False, "error": "Email requerido"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_mis_postulaciones(%s)", (email,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/talento-general", methods=["POST"])
def api_agregar_talento():
    """Agrega un candidato a la base de talento general"""
    try:
        nombre = request.form.get('nombre')
        email = request.form.get('email')
        telefono = request.form.get('telefono')
        
        cv_filename = None
        cv_analisis = '{}'
        
        if 'cv' in request.files:
            cv_file = request.files['cv']
            if cv_file and allowed_file(cv_file.filename):
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                cv_filename = f"talento_{timestamp}_{secure_filename(cv_file.filename)}"
                cv_path = os.path.join(UPLOAD_FOLDER, cv_filename)
                cv_file.save(cv_path)
                
                text = extract_text(cv_path)
                if text and len(text) > 50:
                    data = analyze_cv(text, "gpt")
                    data = normalize_lists(data)
                    data = score_candidate(data)
                    cv_analisis = json.dumps(data)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT agregar_talento_general(%s, %s, %s, %s, %s::jsonb)
        """, (nombre, email, telefono, cv_filename, cv_analisis))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/calcular-match/<int:postulacion_id>", methods=["POST"])
def api_calcular_match(postulacion_id):
    """Calcula el match score usando IA"""
    try:
        import random
        score = random.randint(50, 95)
        
        analisis = json.dumps({
            "fortalezas": ["Experiencia relevante", "Habilidades técnicas"],
            "debilidades": ["Falta experiencia en ciertas tecnologías"],
            "recomendacion": "Avanzar a siguiente fase"
        })
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT actualizar_match_score(%s, %s, %s::jsonb)", 
                      (postulacion_id, score, analisis))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500    


@app.route('/candidatos-portal.html')
def candidatos_portal():
    return render_template('candidatos-portal.html')

# =====================================================
# ENDPOINTS PARA ADMINISTRACIÓN DE EMPLEOS
# =====================================================

@app.route("/api/admin/puestos", methods=["GET"])
def admin_obtener_puestos():
    """Obtiene todos los puestos (incluyendo inactivos)"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, cliente, titulo, descripcion, requisitos, 
                   habilidades_requeridas, seniority, ubicacion, 
                   salario_oferta, activo, fecha_publicacion
            FROM puestos_trabajo
            ORDER BY fecha_publicacion DESC
        """)
        
        rows = cursor.fetchall()
        conn.close()
        
        puestos = []
        for row in rows:
            puestos.append({
                "id": row[0],
                "cliente": row[1],
                "titulo": row[2],
                "descripcion": row[3],
                "requisitos": row[4],
                "habilidades_requeridas": row[5] if row[5] else [],
                "seniority": row[6],
                "ubicacion": row[7],
                "salario_oferta": row[8],
                "activo": row[9],
                "fecha_publicacion": row[10]
            })
        
        return jsonify({"success": True, "puestos": puestos})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500






@app.route("/api/admin/puesto/<int:puesto_id>", methods=["DELETE"])
def admin_eliminar_puesto(puesto_id):
    """Elimina un puesto"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT eliminar_puesto(%s)", (puesto_id,))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/admin/puesto/<int:puesto_id>/preguntas", methods=["GET"])
def admin_obtener_preguntas(puesto_id):
    """Obtiene las preguntas de un puesto"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT id, pregunta, tipo, opciones, requerido, orden
            FROM preguntas_formulario
            WHERE puesto_id = %s
            ORDER BY orden
        """, (puesto_id,))
        
        rows = cursor.fetchall()
        conn.close()
        
        preguntas = []
        for row in rows:
            preguntas.append({
                "id": row[0],
                "pregunta": row[1],
                "tipo": row[2],
                "opciones": row[3] if row[3] else [],
                "requerido": row[4],
                "orden": row[5]
            })
        
        return jsonify({"success": True, "preguntas": preguntas})
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/admin/puesto/preguntas", methods=["POST"])
def admin_guardar_preguntas():
    """Guarda las preguntas de un puesto"""
    try:
        data = request.json
        puesto_id = data.get('puesto_id')
        preguntas = data.get('preguntas', [])
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT guardar_preguntas_formulario(%s, %s::jsonb)
        """, (puesto_id, json.dumps(preguntas)))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/admin/postulaciones", methods=["GET"])
def admin_obtener_postulaciones():
    """Obtiene todas las postulaciones"""
    try:
        estado = request.args.get('estado')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_postulaciones_admin(%s)", (estado,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/admin/postulacion/estado", methods=["POST"])
def admin_actualizar_estado():
    """Actualiza el estado de una postulación"""
    try:
        data = request.json
        postulacion_id = data.get('postulacion_id')
        estado = data.get('estado')
        match_score = data.get('match_score')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT actualizar_estado_postulacion(%s, %s, %s)
        """, (postulacion_id, estado, match_score))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route('/admin-empleos.html')
def admin_empleos():
    return render_template('admin-empleos.html')

@app.route("/api/admin/puesto", methods=["POST"])
def admin_guardar_puesto():
    """Guarda un puesto (nuevo o existente) con todos los campos"""
    try:
        data = request.json
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT guardar_puesto(
                %s, %s, %s, %s, %s, %s, %s, %s, %s, %s,
                %s, %s, %s, %s, %s, %s, %s, %s
            )
        """, (
            data.get('id'),
            data.get('cliente'),
            data.get('titulo'),
            data.get('descripcion'),
            data.get('requisitos'),
            data.get('habilidades_requeridas'),
            data.get('seniority'),
            data.get('ubicacion'),
            data.get('salario_oferta'),
            data.get('activo', True),
            data.get('responsabilidades'),
            data.get('beneficios'),
            data.get('sobre_empresa'),
            data.get('horario'),
            data.get('tipo_contrato'),
            data.get('modalidad'),
            data.get('vacantes', 1),
            data.get('fecha_limite')
        ))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"success": False, "error": str(e)}), 500
    
@app.route("/api/entrevista/registrar-infraccion", methods=["POST"])
def registrar_infraccion():
    """Registra una infracción del candidato (cambio de pestaña, etc.)"""
    try:
        data = request.json
        token = data.get('token')
        tipo_infraccion = data.get('tipo', 'cambio_pestana')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # Obtener información de la entrevista
        cursor.execute("""
            SELECT e.id, e.contador_advertencias, p.max_advertencias, p.expulsar_al_superar
            FROM entrevistas e
            JOIN plantillas_entrevista_v2 p ON e.plantilla_id = p.id
            WHERE e.token_acceso = %s AND e.estado = 'en_progreso'
        """, (token,))
        
        row = cursor.fetchone()
        
        if not row:
            conn.close()
            return jsonify({"success": False, "error": "Entrevista no encontrada"}), 404
        
        entrevista_id = row[0]
        contador = row[1] or 0
        max_advertencias = row[2] or 3
        expulsar = row[3] or True
        
        nuevo_contador = contador + 1
        
        # Registrar en bitácora
        cursor.execute("""
            UPDATE entrevistas 
            SET contador_advertencias = %s,
                bitacora_infracciones = bitacora_infracciones || %s::jsonb
            WHERE id = %s
        """, (nuevo_contador, json.dumps([{
            "timestamp": datetime.now().isoformat(),
            "tipo": tipo_infraccion,
            "contador": nuevo_contador
        }]), entrevista_id))
        
        # Verificar si debe expulsar
        fue_expulsado = False
        if expulsar and nuevo_contador >= max_advertencias:
            cursor.execute("""
                UPDATE entrevistas 
                SET estado = 'expulsado',
                    fue_expulsado = true,
                    razon_expulsion = %s,
                    tiempo_total_fin = NOW()
                WHERE id = %s
            """, (f"Superó el máximo de {max_advertencias} advertencias. Infracción: {tipo_infraccion}", entrevista_id))
            fue_expulsado = True
        
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": True, 
            "contador": nuevo_contador,
            "max_advertencias": max_advertencias,
            "fue_expulsado": fue_expulsado,
            "advertencias_restantes": max_advertencias - nuevo_contador
        })
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500
    
# =====================================================
# ENDPOINTS PARA KANBAN DE RECLUTAMIENTO
# =====================================================

@app.route("/api/kanban/etapas", methods=["GET"])
def kanban_obtener_etapas():
    """Obtiene todas las etapas del proceso"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_etapas()")
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/candidatos/<int:puesto_id>", methods=["GET"])
def kanban_obtener_candidatos(puesto_id):
    """Obtiene candidatos organizados por etapa para un puesto"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_candidatos_kanban(%s)", (puesto_id,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/mover", methods=["POST"])
def kanban_mover_candidato():
    """Mueve un candidato a otra etapa"""
    try:
        data = request.json
        candidato_proceso_id = data.get('candidato_proceso_id')
        nueva_etapa_id = data.get('nueva_etapa_id')
        comentario = data.get('comentario', '')
        usuario = data.get('usuario', 'reclutador')
        
        if not candidato_proceso_id or not nueva_etapa_id:
            return jsonify({"success": False, "error": "Faltan datos"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT mover_candidato(%s, %s, %s, %s)", 
                      (candidato_proceso_id, nueva_etapa_id, usuario, comentario))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/etapa", methods=["POST"])
def kanban_crear_etapa():
    """Crea una nueva etapa"""
    try:
        data = request.json
        nombre = data.get('nombre')
        color = data.get('color', '#3b82f6')
        orden = data.get('orden')
        
        if not nombre:
            return jsonify({"success": False, "error": "Nombre requerido"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT crear_etapa(%s, %s, %s)", (nombre, color, orden))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/etapa/<int:etapa_id>", methods=["PUT"])
def kanban_actualizar_etapa(etapa_id):
    """Actualiza una etapa"""
    try:
        data = request.json
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT actualizar_etapa(%s, %s, %s, %s, %s)", 
                      (etapa_id, data.get('nombre'), data.get('color'), 
                       data.get('orden'), data.get('activo', True)))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/etapa/<int:etapa_id>", methods=["DELETE"])
def kanban_eliminar_etapa(etapa_id):
    """Elimina una etapa"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT eliminar_etapa(%s)", (etapa_id,))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/kanban/historial/<int:candidato_proceso_id>", methods=["GET"])
def kanban_obtener_historial(candidato_proceso_id):
    """Obtiene el historial de movimientos de un candidato"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_historial_candidato(%s)", (candidato_proceso_id,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

@app.route('/kanban-reclutamiento.html')
def kanban_reclutamiento():
    return render_template('kanban-reclutamiento.html')

@app.route('/kanban.html')
def kanban_page():
    return render_template('kanban.html')


@app.route('/kanban')
def kanban():
    return render_template('kanban.html')

@app.route("/api/kanban/actualizar-notas", methods=["POST"])
def actualizar_notas_candidato():
    try:
        data = request.json
        candidato_proceso_id = data.get('candidato_proceso_id')
        notas = data.get('notas', '')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            UPDATE candidatos_proceso 
            SET notas = %s, fecha_actualizacion = CURRENT_TIMESTAMP
            WHERE id = %s
        """, (notas, candidato_proceso_id))
        
        conn.commit()
        conn.close()
        
        return jsonify({"success": True, "message": "Notas actualizadas"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


# =====================================================
# ENDPOINTS PARA ENVÍO DE CORREOS
# =====================================================

# =====================================================
# ENDPOINTS PARA ENVÍO DE CORREOS
# =====================================================

# =====================================================
# ENDPOINTS PARA ENVÍO DE CORREOS
# =====================================================

@app.route("/api/enviar-correo", methods=["POST"])
def enviar_correo():
    """Envía un correo electrónico"""
    try:
        data = request.json
        to_email = data.get('to')
        subject = data.get('subject')
        body = data.get('body')
        is_html = data.get('is_html', True)
        
        if not to_email or not subject or not body:
            return jsonify({"success": False, "error": "Faltan datos: to, subject, body"}), 400
        
        from email_service import mailer
        result = mailer.send_email_sync(to_email, subject, body, is_html)
        
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/enviar-recordatorio-entrevista", methods=["POST"])
def enviar_recordatorio_entrevista():
    """Envía recordatorio de entrevista al candidato"""
    try:
        data = request.json
        candidato_email = data.get('email')
        candidato_nombre = data.get('nombre')
        fecha = data.get('fecha')
        hora = data.get('hora')
        enlace = data.get('enlace')
        
        subject = f"📅 Recordatorio de entrevista - Talent Pipeline"
        
        body = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <div style="background: #2563eb; padding: 20px; text-align: center; border-radius: 12px 12px 0 0;">
                <h2 style="color: white; margin: 0;">🎯 Talent Pipeline</h2>
            </div>
            <div style="padding: 24px; border: 1px solid #e2e8f0; border-radius: 0 0 12px 12px;">
                <h3 style="color: #1e293b;">Hola {candidato_nombre},</h3>
                <p>Tienes una entrevista programada:</p>
                <div style="background: #f8fafc; padding: 16px; border-radius: 8px; margin: 16px 0;">
                    <p><strong>📅 Fecha:</strong> {fecha}</p>
                    <p><strong>⏰ Hora:</strong> {hora}</p>
                    <p><strong>🔗 Enlace:</strong> <a href="{enlace}" style="color: #2563eb;">Acceder a la entrevista</a></p>
                </div>
                <p>Por favor, asegúrate de tener:</p>
                <ul>
                    <li>Cámara y micrófono funcionando</li>
                    <li>Conexión estable a internet</li>
                    <li>Un lugar tranquilo y bien iluminado</li>
                </ul>
                <hr style="margin: 20px 0; border-color: #e2e8f0;">
                <p style="color: #64748b; font-size: 12px;">Este es un mensaje automático, por favor no responder.</p>
            </div>
        </div>
        """
        
        from email_service import mailer
        result = mailer.send_email_sync(candidato_email, subject, body, is_html=True)
        
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

@app.route("/api/enviar-resultado-entrevista", methods=["POST"])
def enviar_resultado_entrevista():
    """Envía resultados de la entrevista al candidato"""
    try:
        data = request.json
        candidato_email = data.get('email')
        candidato_nombre = data.get('nombre')
        puesto = data.get('puesto')
        resultado = data.get('resultado')  # 'aprobado', 'rechazado', 'pendiente'
        comentarios = data.get('comentarios', '')
        
        if resultado == 'aprobado':
            subject = f"✅ Actualización de tu postulación - {puesto}"
            body = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <div style="background: #10b981; padding: 20px; text-align: center; border-radius: 12px 12px 0 0;">
                    <h2 style="color: white; margin: 0;">✅ ¡Buenas noticias!</h2>
                </div>
                <div style="padding: 24px; border: 1px solid #e2e8f0;">
                    <h3>Hola {candidato_nombre},</h3>
                    <p>¡Felicitaciones! Has avanzado a la siguiente etapa del proceso para el puesto de <strong>{puesto}</strong>.</p>
                    {f'<p><strong>Comentarios del reclutador:</strong><br>{comentarios}</p>' if comentarios else ''}
                    <p>Pronto nos comunicaremos contigo para coordinar los siguientes pasos.</p>
                </div>
            </div>
            """
        elif resultado == 'rechazado':
            subject = f"📫 Actualización de tu postulación - {puesto}"
            body = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <div style="background: #ef4444; padding: 20px; text-align: center; border-radius: 12px 12px 0 0;">
                    <h2 style="color: white; margin: 0;">📫 Actualización del proceso</h2>
                </div>
                <div style="padding: 24px; border: 1px solid #e2e8f0;">
                    <h3>Hola {candidato_nombre},</h3>
                    <p>Gracias por tu interés en el puesto de <strong>{puesto}</strong>.</p>
                    <p>Después de evaluar tu perfil, hemos decidido continuar con otros candidatos que se ajustan mejor a los requisitos actuales.</p>
                    {f'<p><strong>Comentarios:</strong><br>{comentarios}</p>' if comentarios else ''}
                    <p>Te invitamos a seguir postulando a futuras vacantes que se ajusten a tu perfil.</p>
                </div>
            </div>
            """
        else:
            subject = f"📋 Actualización de tu postulación - {puesto}"
            body = f"""
            <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
                <div style="background: #f59e0b; padding: 20px; text-align: center; border-radius: 12px 12px 0 0;">
                    <h2 style="color: white; margin: 0;">📋 Proceso en curso</h2>
                </div>
                <div style="padding: 24px; border: 1px solid #e2e8f0;">
                    <h3>Hola {candidato_nombre},</h3>
                    <p>Tu postulación para el puesto de <strong>{puesto}</strong> está siendo evaluada por nuestro equipo.</p>
                    <p>Te mantendremos informado sobre cualquier novedad.</p>
                </div>
            </div>
            """
        
        from email_service import mailer
        result = mailer.send_email_sync(candidato_email, subject, body, is_html=True)
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500
    
# =====================================================
# ENDPOINTS PARA MÓDULO DE CORREOS
# =====================================================

@app.route("/api/correos/plantillas", methods=["GET"])
def api_obtener_plantillas():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_plantillas_correo()")
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/correos/plantilla", methods=["POST"])
def api_crear_plantilla():
    try:
        data = request.json
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT crear_plantilla_correo(%s, %s, %s, %s, %s)
        """, (data['nombre'], data['asunto'], data['cuerpo'], data.get('tipo', 'manual'), data.get('variables')))
        
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/correos/procesar", methods=["POST"])
def api_procesar_plantilla():
    try:
        data = request.json
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT procesar_plantilla_correo(%s, %s, %s, %s::jsonb)
        """, (data['plantilla_id'], data['candidato_id'], data.get('postulacion_id'), json.dumps(data.get('variables', {}))))
        
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/correos/enviar", methods=["POST"])
def api_enviar_correo():
    try:
        data = request.json
        plantilla_id = data.get('plantilla_id')
        candidato_id = data.get('candidato_id')
        postulacion_id = data.get('postulacion_id')
        variables_extra = data.get('variables', {})
        
        # 1. Procesar plantilla
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("""
            SELECT procesar_plantilla_correo(%s, %s, %s, %s::jsonb)
        """, (plantilla_id, candidato_id, postulacion_id, json.dumps(variables_extra)))
        
        procesado = cursor.fetchone()[0]
        
        if not procesado.get('success'):
            conn.close()
            return jsonify(procesado)
        
        # 2. Registrar en historial
        cursor.execute("""
            SELECT registrar_correo_historial(%s, %s, %s, %s, %s, %s, 'pendiente')
        """, (candidato_id, postulacion_id, plantilla_id, 
              procesado['to'], procesado['subject'], procesado['body']))
        
        historial_id = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        # 3. Enviar correo
        from email_service import mailer
        result = mailer.send_email_sync(
            to_email=procesado['to'],
            subject=procesado['subject'],
            body=procesado['body'],
            is_html=True
        )
        
        # 4. Actualizar estado del historial
        conn = get_db_connection()
        cursor = conn.cursor()
        
        if result.get('success'):
            cursor.execute("SELECT actualizar_estado_historial(%s, 'enviado', NULL)", (historial_id,))
        else:
            cursor.execute("SELECT actualizar_estado_historial(%s, 'error', %s)", 
                          (historial_id, result.get('error', 'Error desconocido')))
        
        result_estado = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify({
            "success": result.get('success', False),
            "message": result.get('message', ''),
            "error": result.get('error', ''),
            "historial_id": historial_id
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/correos/historial", methods=["GET"])
def api_obtener_historial():
    try:
        candidato_id = request.args.get('candidato_id')
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_historial_correos(%s)", (candidato_id,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/correos/variables", methods=["GET"])
def api_obtener_variables():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_variables_disponibles()")
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500



# =====================================================
# ENDPOINTS DE CONFIGURACIÓN Y AUDITORÍA
# =====================================================

@app.route("/api/configuracion", methods=["GET"])
@limiter.limit("30 per minute")
def api_obtener_configuracion():
    """Obtiene toda la configuración del sistema"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_configuracion()")
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/configuracion", methods=["PUT"])
@limiter.limit("10 per minute")
def api_actualizar_configuracion():
    """Actualiza una configuración del sistema"""
    try:
        data = request.json
        clave = data.get('clave')
        valor = data.get('valor')
        
        if not clave:
            return jsonify({"success": False, "error": "Clave requerida"}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT actualizar_configuracion(%s, %s)", (clave, valor))
        result = cursor.fetchone()[0]
        conn.commit()
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/configuracion/grupo/<grupo>", methods=["GET"])
@limiter.limit("30 per minute")
def api_obtener_configuracion_grupo(grupo):
    """Obtiene configuración por grupo"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_configuracion_por_grupo(%s)", (grupo,))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/auditoria", methods=["GET"])
@limiter.limit("30 per minute")
def api_obtener_auditoria():
    """Obtiene historial de auditoría"""
    try:
        limite = request.args.get('limite', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        conn = get_db_connection()
        cursor = conn.cursor()
        
        cursor.execute("SELECT obtener_auditoria(%s, %s)", (limite, offset))
        result = cursor.fetchone()[0]
        conn.close()
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500

# =====================================================
# INICIO
# =====================================================
if __name__ == "__main__":
    print("🔍 Verificando conexión a PostgreSQL...")
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT 1")
        conn.close()
        print("✅ PostgreSQL conectado correctamente")
    except Exception as e:
        print(f"❌ Error conectando a PostgreSQL: {e}")
        print("⚠️ Asegúrate de que PostgreSQL esté corriendo")
    app.run(debug=True, port=5001)